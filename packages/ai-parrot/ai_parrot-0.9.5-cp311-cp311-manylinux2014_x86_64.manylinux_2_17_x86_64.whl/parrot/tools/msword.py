from datetime import datetime
from typing import Any, Dict, List, Optional, Type, Union
import re
import tempfile
import os
from pathlib import Path
import uuid
import asyncio
import traceback
import logging
from urllib.parse import urlparse
import aiohttp
import aiofiles
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel, Field, ConfigDict
from langchain.tools import BaseTool
from markdownify import markdownify as md
import mammoth
import markdown
from bs4 import BeautifulSoup, NavigableString
import html2text
from navconfig import BASE_DIR


logging.getLogger('MARKDOWN').setLevel(logging.ERROR)  # Suppress markdown warnings


class DocxInput(BaseModel):
    """
    Input schema for the DocxGeneratorTool.  Users can supply:
    • text (required): the transcript or Markdown to saved as a DOCX File.
    • output_filename: (Optional) a custom filename (including extension) for the generated DOCX.
    • template_name: (Optional) name of the HTML template (e.g. 'report.html') to render
    • template_vars: (Optional) dict of variables to pass into the template (e.g. title, author, date)
    • stylesheets: (Optional) list of CSS file paths (relative to your templates dir) to apply
    • output_dir: (Optional) directory where the DOCX file will be saved.
    • docx_template: (Optional) path to a DOCX template file to use as base
    """
    # Add a model_config to prevent additional properties
    model_config = ConfigDict(extra='forbid')

    text: str = Field(..., description="The text (plaintext or Markdown) to convert to DOCX File")
    # If you’d like users to control the output filename/location:
    output_filename: Optional[str] = Field(
        None,
        description="(Optional) A custom filename (including extension) for the generated DOCX."
    )
    template_name: Optional[str] = Field(
        None,
        description="Name of the HTML template (e.g. 'report.html') to render"
    )
    template_vars: Optional[Dict[str, str]] = Field(
        None,
        description="Dict of variables to pass into the template (e.g. title, author, date)"
    )
    stylesheets: Optional[List[str]] = Field(
        None,
        description="List of CSS file paths (relative to your templates dir) to apply"
    )
    output_dir: Optional[str] = Field(
        None,
        description="Directory where the DOCX file will be saved."
    )
    docx_template: Optional[str] = Field(
        None,
        description="Path to a DOCX template file to use as base document"
    )


class DocxGeneratorTool(BaseTool):
    """Microsoft Word DOCX Generator Tool."""
    name: str = "generate_ms_word_document"
    description: str = "Use this tool for generating DOCX, provide text in markdown or HTML format with sections, headings."
    output_dir: Optional[Path] = None
    env: Optional[Environment] = None
    templates_dir: Optional[Path] = None

    # Add a proper args_schema for tool-calling compatibility
    args_schema: Type[BaseModel] = DocxInput

    def __init__(
        self,
        templates_dir: Path = None,
        output_dir: str = None
    ):
        """Initialize the DOCX generator tool."""
        super().__init__()
        self.output_dir = Path(output_dir) if output_dir else BASE_DIR.joinpath("static", "documents", "docs")
        if not self.output_dir.exists():
            self.output_dir.mkdir(parents=True, exist_ok=True)
        # Initialize Jinja2 environment for HTML templates
        if templates_dir:
            self.templates_dir = templates_dir
            self.env = Environment(
                loader=FileSystemLoader(str(templates_dir)),
                autoescape=True
            )
        self.templates_dir = templates_dir


    async def _arun(
        self,
        text: str,
        output_filename: Optional[str] = None,
        **kwargs: Any
    ) -> Dict[str, Any]:
        """
        LangChain will call this with keyword args matching DocxInput, e.g.:
        _arun(text="Hello", output_dir="documents/", …)
        """
        try:
            # 1) Build a dict of everything LangChain passed us
            payload_dict = {
                "text": text,
                "output_filename": output_filename,
                **kwargs
            }
            # 2) Let Pydantic validate & coerce
            payload = DocxInput(**{k: v for k, v in payload_dict.items() if v is not None})
            # 3) Call the “real” generator
            return await self._generate_docx(payload)
        except Exception as e:
            print(f"❌ Error in DocxGeneratorTool._arun: {e}")
            print(traceback.format_exc())
            return {"error": str(e)}

    def _run(
        self,
        text: Union[str, Dict[str, Any]],
        **kwargs
    ) -> Dict[str, Any]:
        """
        Synchronous entrypoint. If text_or_json is a JSON string, we load it first.
        Otherwise, assume it’s already a dict of the correct shape.
        """
        try:
            # Validate with DocxInput
            payload = DocxInput(text=text, **kwargs)
        except Exception as e:
            return {"error": f"Invalid input: {e}"}
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                return loop.run_until_complete(self._generate_docx(payload))
            else:
                return asyncio.run(self._generate_docx(payload))
        except RuntimeError:
            # If the event loop is not running, we can run it in a new loop
            return asyncio.run(self._generate_docx(payload))

    def _render_template(self, input_data: DocxInput) -> str:
        """Render text through Jinja2 template if provided."""
        if not input_data.template_name or not self.env:
            return input_data.text

        try:
            template = self.env.get_template(input_data.template_name)
            template_vars = input_data.template_vars or {}

            # Add some default variables
            template_vars.setdefault('content', input_data.text)
            template_vars.setdefault('date', datetime.now().strftime('%Y-%m-%d'))
            template_vars.setdefault('timestamp', datetime.now().isoformat())

            return template.render(**template_vars)
        except Exception as e:
            print(f"Template rendering failed: {e}")
            return input_data.text

    def _detect_content_type(self, text: str) -> str:
        """Detect if content is HTML, Markdown, or plain text."""
        text_stripped = text.strip()

        # Simple HTML detection
        if (text_stripped.startswith('<') and text_stripped.endswith('>')) or \
            ('<html' in text_stripped.lower() or '<div' in text_stripped.lower() or
            '<p' in text_stripped.lower() or '<h1' in text_stripped.lower()):
            return 'html'

        # Markdown detection (look for common markdown patterns)
        markdown_patterns = [
            r'^#{1,6}\s',  # Headers
            r'^\*\s',      # Bullet points
            r'^\d+\.\s',   # Numbered lists
            r'\*\*.*?\*\*', # Bold
            r'\*.*?\*',    # Italic
            r'`.*?`',      # Code
            r'\[.*?\]\(.*?\)', # Links
        ]

        for pattern in markdown_patterns:
            if re.search(pattern, text_stripped, re.MULTILINE):
                return 'markdown'

        return 'markdown'  # Default to markdown for processing

    async def _generate_docx(self, payload: DocxInput) -> dict:
        """Generate a DOCX document from markdown text."""
        # Process the text through Jinja2 template if provided
        processed_text = self._render_template(payload)
        # Detect content type and convert to DOCX
        content_type = self._detect_content_type(processed_text)
        # Create or load DOCX document
        doc = self._create_document(payload.docx_template)

        # Convert content to DOCX based on type
        if content_type == 'html':
            self._html_to_docx(processed_text, doc)
        else:  # markdown or plain text
            processed_text = self._preprocess_markdown(processed_text)
            # Convert markdown to HTML first
            html_content = self._markdown_to_html(processed_text)
            self._html_to_docx(html_content, doc)
        try:
            # Generate filename and save
            output_path = self._save_document(doc, payload)
            return {
                "status": "success",
                "file_path": str(output_path),
                "filename": output_path.name,
                "message": f"DOCX file successfully created at {output_path}"
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "message": f"Failed to create DOCX file: {str(e)}"
            }

    def _preprocess_markdown(self, text):
        """Preprocess markdown to handle common issues."""
        # Replace placeholder variables with empty strings
        text = re.sub(r'\{[a-zA-Z0-9_]+\}', '', text)

        # Handle f-strings that weren't evaluated
        text = re.sub(r'f"""(.*?)"""', r'\1', text, flags=re.DOTALL)
        text = re.sub(r"f'''(.*?)'''", r'\1', text, flags=re.DOTALL)

        # Remove triple backticks and language indicators (common in code blocks)
        text = re.sub(r'```[a-zA-Z]*\n', '', text)
        text = re.sub(r'```', '', text)

        # Fix any heading issues (ensure space after #)
        text = re.sub(r'(#+)([^ \n])', r'\1 \2', text)

        return text

    def _markdown_to_html(self, markdown_text: str) -> str:
        """Convert markdown to HTML."""
        try:
            html = markdown.markdown(
                markdown_text,
                extensions=['extra', 'codehilite', 'toc']
            )
            return html
        except Exception as e:
            print(f"Markdown conversion failed: {e}")
            # Fallback: wrap in paragraphs
            paragraphs = markdown_text.split('\n\n')
            html_paragraphs = [f'<p>{p.replace(chr(10), "<br>")}</p>' for p in paragraphs if p.strip()]
            return '\n'.join(html_paragraphs)

    def _create_document(self, template_path: Optional[str] = None) -> Document:
        """Create or load DOCX document."""
        if template_path and Path(template_path).exists():
            return Document(template_path)

        # Create new document with basic styling
        doc = Document()

        # Set up styles
        self._setup_document_styles(doc)

        return doc

    def _setup_document_styles(self, doc: Document):
        """Set up basic document styles."""
        try:
            styles = doc.styles

            # Normal style
            normal = styles['Normal']
            normal.font.name = 'Calibri'
            normal.font.size = Pt(11)

            # Heading styles
            for i in range(1, 7):
                heading_name = f'Heading {i}'
                if heading_name in styles:
                    heading = styles[heading_name]
                    heading.font.name = 'Calibri'
                    heading.font.size = Pt(16 - i)

        except Exception as e:
            print(f"Style setup failed: {e}")

    def _html_to_docx(self, html_content: str, doc: Document):
        """Convert HTML content to DOCX document."""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')

            # Process each element in the HTML
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'ul', 'ol', 'li', 'br']):
                self._process_html_element(element, doc)

        except Exception as e:
            print(f"HTML to DOCX conversion failed: {e}")
            # Fallback: add as plain text
            doc.add_paragraph(html_content)

    def _process_html_element(self, element, doc: Document):
        """Process individual HTML elements."""
        tag_name = element.name.lower()

        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag_name[1])
            heading = doc.add_heading(self._get_text_content(element), level=level)

        elif tag_name in ['p', 'div']:
            text = self._get_text_content(element)
            if text.strip():
                paragraph = doc.add_paragraph()
                self._add_formatted_text(paragraph, element)

        elif tag_name == 'table':
            self._process_table(element, doc)

        elif tag_name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = self._get_text_content(li)
                if text.strip():
                    doc.add_paragraph(text, style='List Bullet' if tag_name == 'ul' else 'List Number')

        elif tag_name == 'br':
            doc.add_paragraph()

    def _get_text_content(self, element) -> str:
        """Extract text content from HTML element."""
        if isinstance(element, NavigableString):
            return str(element)

        text_parts = []
        for content in element.contents:
            if isinstance(content, NavigableString):
                text_parts.append(str(content))
            else:
                text_parts.append(self._get_text_content(content))

        return ''.join(text_parts).strip()

    def _process_table(self, table_element, doc: Document):
        """Process HTML table and convert to DOCX table."""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Create table with appropriate dimensions
        cols = len(rows[0].find_all(['td', 'th']))
        table = doc.add_table(rows=0, cols=cols)
        table.style = 'Table Grid'

        for row in rows:
            cells = row.find_all(['td', 'th'])
            table_row = table.add_row()
            for i, cell in enumerate(cells):
                if i < len(table_row.cells):
                    table_row.cells[i].text = self._get_text_content(cell)

    def _add_formatted_text(self, paragraph, element):
        """Add formatted text to paragraph maintaining basic formatting."""
        if isinstance(element, NavigableString):
            paragraph.add_run(str(element))
            return

        for content in element.contents:
            if isinstance(content, NavigableString):
                run = paragraph.add_run(str(content))
            else:
                run = paragraph.add_run(self._get_text_content(content))

                # Apply basic formatting
                if content.name == 'strong' or content.name == 'b':
                    run.bold = True
                elif content.name == 'em' or content.name == 'i':
                    run.italic = True
                elif content.name == 'code':
                    run.font.name = 'Courier New'

    def _save_document(self, doc: Document, input_data: DocxInput) -> Path:
        """Save the document and return the file path."""
        # Determine output directory
        output_dir = Path(input_data.output_dir) if input_data.output_dir else self.output_dir
        output_dir.mkdir(parents=True, exist_ok=True)

        # Generate filename
        if input_data.output_filename:
            filename = input_data.output_filename
            if not filename.endswith('.docx'):
                filename += '.docx'
        else:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"document_{timestamp}_{uuid.uuid4().hex[:8]}.docx"

        output_path = output_dir / filename

        # Save the document
        doc.save(str(output_path))

        return output_path

class WordToMarkdownTool(BaseTool):
    """Converts a Word document to Markdown format by downloading it from a URL."""
    name: str = "word_to_markdown_tool"
    description: str = (
        "Converts a Word document to Markdown format from a URL. "
        "This tool downloads the Word document from the provided URL, "
        "converts it to Markdown format, and returns the content. "
        "Useful for processing Word documents and making them easier to analyze by LLMs."
        "\nThe input must be the complete URL of the Word document."
    )
    return_direct: bool = False
    _temp_dir: Optional[str] = None

    async def _download_file(self, url: str) -> str:
        """Downloads a file from a URL to a temporary file."""
        # Create a temporary directory if it doesn't exist
        if not self._temp_dir:
            self._temp_dir = tempfile.mkdtemp()

        # Get the filename from the URL
        parsed_url = urlparse(url)
        filename = os.path.basename(parsed_url.path)
        if not filename.endswith(('.docx', '.doc')):
            filename += '.docx'  # Add extension if it doesn't exist

        # Complete path to the temporary file
        file_path = os.path.join(self._temp_dir, filename)

        # Download the file
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise Exception(f"Error downloading the file: {response.status}")

                # Save the file
                async with aiofiles.open(file_path, 'wb') as f:
                    await f.write(await response.read())

        return file_path

    async def _convert_to_markdown(self, file_path: str) -> str:
        """Converts a Word document to Markdown."""
        # Use mammoth to convert to HTML and then to Markdown
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            markdown_text = md(html)

            # If there are warning messages, add them as a comment at the beginning
            if result.messages:
                warnings = "\n".join([f"<!-- Warning: {msg} -->" for msg in result.messages])
                markdown_text = f"{warnings}\n\n{markdown_text}"

            return markdown_text

    async def _process_word_document(self, url: str) -> Dict[str, Any]:
        """Processes a Word document from a URL and converts it to Markdown."""
        try:
            file_path = await self._download_file(url)
            markdown_text = await self._convert_to_markdown(file_path)

            # Cleanup of temporary files
            if os.path.exists(file_path):
                os.remove(file_path)

            return {
                "markdown": markdown_text,
                "source_url": url,
                "success": True
            }
        except Exception as e:
            return {
                "error": str(e),
                "source_url": url,
                "success": False
            }
        finally:
            # Ensure cleanup of the temporary directory if it's empty
            if self._temp_dir and os.path.exists(self._temp_dir) and not os.listdir(self._temp_dir):
                os.rmdir(self._temp_dir)

    async def _arun(self, url: str) -> Dict[str, Any]:
        """Runs the tool asynchronously."""
        return await self._process_word_document(url)

    def _run(self, url: str) -> Dict[str, Any]:
        """Runs the tool synchronously."""
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                return loop.run_until_complete(self._process_word_document(url))
            else:
                return asyncio.run(self._process_word_document(url))
        except RuntimeError:
            # If the event loop is not running, we can run it in a new loop
            return asyncio.run(self._process_word_document(url))

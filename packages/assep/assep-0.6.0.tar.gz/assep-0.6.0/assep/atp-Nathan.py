import matplotlib.pyplot as plt
import scipy.io as sio
import subprocess
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
import time
import os
import shutil

class ATP:
    def __init__(self, path_fileATP):
        # Diretórios padrão
        self.path_bat = r"C:\ATP\tools\runATP.bat"
        self.path_mat = r"C:\ATP\Pl42mat09\Pl42mat.exe"
        self.path_word_template_report = r"C:\ATP\Template_report_empty.docx"
        self.path_fileATP = path_fileATP

        # Diretórios relativos
        self.folder_atp = path_fileATP[:path_fileATP.rfind("\\")+1]
        self.output_file = self.folder_atp + "resultado_estatistico.txt"
        self.word_file = self.folder_atp + "\\curvas.docx"

        # Parâmetros de simulação
        self.num_simul = 200                            # Número de simulações
        self.number_curves_plot = 5                     # Número de curvas a plotar por tipo de grandeza
        self.curvas_sempre_plotar = ["TF", "TF_D"]      # Curvas para sempre plotar, independente do resultado
        self.curvas_sempre_plotar = ["CBA52", "CBA52_TCAT_",
                                     "NTF1","NTF1_R","PR"]      # Curvas para sempre plotar, independente do resultado
        self.executar_ATP_simul = True                  # Roda o  estudo estatístico via ATP
        self.run_atp_shot = True                        # Roda o conversor PL4 para MAT
        self.apply_envoltoria = True                    # Plotar a envoltória de tensão
        self.fig_widht = 7                              # Tamanho das figuras no relatório Word
        self.time_simulation = 1.0                      # Tempo de simulação
        self.time_simulation_shot = 5.0                 # Tempo de simulação do shot
        self.dpi_res = 500                              # Qualidade png   
        self.slice_df_visualiza = 0.10                  # Filtro os 5% inferiores e superiores na hora do ZOOM
        self.open_word = False

        # Diretórios CSVs
        self.path_csv_dfsts = self.folder_atp+"df_sts.csv"
        self.path_csv_dfsts_filt = self.folder_atp+"df_sts_filter.csv"
        self.path_csv_var = self.folder_atp + "df_variables.csv"
        self.path_csv_res = self.folder_atp + "df_report.csv"
        self.path_csv_sht = self.folder_atp + "df_shots.csv"
        self.path_csv_plt = self.folder_atp + "df_plots.csv"
        self.path_docx_report = self.folder_atp + "relatório_TEM.docx"   
        
        self.dic_relacao_regime = {"TF_D": "TF"}
        self.dic_relacoes = {}

    #==============================================================================================
    #
    # FUNC 01
    #
    #==============================================================================================
    def adjust_num_simulations(self, num_simul=""):
        if num_simul != "":
            self.num_simul = num_simul

        with open(self.path_fileATP, 'r') as f:
            lines = f.readlines()

        target_line = 'C  dT  >< Tmax >< Xopt >< Copt ><Epsiln>\n'

        try:
            target_index = lines.index(target_line)
        except ValueError:
            print("Erro modifica número simulações.")
        else:
            # Modify the lines following the target line
            lines[target_index + 2] = lines[target_index + 2][:69] + "{:3d}".format(self.num_simul) + "\n"

            # Write the updated content back to the file
            with open(self.path_fileATP, 'w') as file:
                file.writelines(lines)
        return
    
    def adjust_time_simulations(self, type_file="main", path_shot=""):
        if type_file == "main":
            with open(self.path_fileATP, 'r') as f:
                lines = f.readlines()
            target_line = 'C  dT  >< Tmax >< Xopt >< Copt ><Epsiln>\n'

            try:
                target_index = lines.index(target_line)
            except ValueError:
                print("Erro modifica tempo simulações.")
            else:
                # Modify the lines following the target line
                lines[target_index + 1] = lines[target_index + 1][:8] + "{:8.3f}".format(self.time_simulation) + lines[target_index + 1][16:-1] + "\n"

            # Write the updated content back to the file
            with open(self.path_fileATP, 'w') as file:
                file.writelines(lines)

        elif type_file == "shot":
            with open(path_shot, 'r') as f:
                lines = f.readlines()
            target_line = 'C $DUMMY, XYZ000                                                                \n'

            try:
                target_index = lines.index(target_line)
            except ValueError:
                print("Erro modifica tempo simulações.")
            else:
                # Modify the lines following the target line
                lines[target_index + 1] = lines[target_index + 1][:8] + "{:8.3f}".format(self.time_simulation_shot) + lines[target_index + 1][16:-1] + "\n"

            # Write the updated content back to the file
            with open(path_shot, 'w') as file:
                file.writelines(lines)

        return

    def remove_files_temp(self, folder_atp):
        for filename in os.listdir(folder_atp):
            if filename.endswith(".tmp") or filename.endswith(".bin") or filename.endswith(".dbg") or filename.endswith(".MAT"):
                file_path = os.path.join(folder_atp, filename)
                try:
                    os.remove(file_path)
                except:
                    pass

    def run_ATP(self, path_file="", output_file="", make_backup = True):
        """ 
        Roda o ATP
        """
        if path_file == "":
            path_file = self.path_fileATP
        if output_file == "":
            output_file = self.output_file

        comando = f'{self.path_bat} "{path_file}"'
        with open(output_file, "w+") as f:
            processo = subprocess.Popen(comando, shell=True, stdout=f)

        ultima_linha = "X"
        while ultima_linha[:20] != "Total Execution Time":
            with open(output_file, "r") as f:
                linhas = f.readlines()
                ultima_linha = linhas[-1].strip() if linhas else ''

        folder_atp = self.path_fileATP[:self.path_fileATP.rfind("\\")+1]        
        self.remove_files_temp(folder_atp)

        if make_backup:
            target = output_file[:-4] + ".lis"
            shutil.copyfile(output_file, target)

        return
    
    #==============================================================================================
    #
    # FUNC 02
    #
    #==============================================================================================
    def check_float(self, number):
        try:
            float(number)
            return True
        except ValueError:
            return False
    
    def get_df_sts(self):
        """"
        Monta o dataframe com os dados estatísticos
        """
        ini_bool = False
        list_data = []
        list_dics = []
        with open(self.output_file, 'r') as file:
            linhas = file.readlines()
            for i, linha in enumerate(linhas):
                # Indica que vou iniciar a apresentação dos resultados estatísticos
                if 'Initialize variables as  "FIND"  begins.          |FIND' in linha:
                    ini_bool = True

                # Indica que vou finalizar a coleta dos resultados estatísticos
                if 'Exit  FIND  by resuming statistical tabulations.  |QUIT' in linha:
                    ini_bool = False
                    
                    # Coleta de Valores
                    dic_sts = {}
                    dic_sts["grandeza"] = (list_data[2][22:39]).strip()
                    dic_sts["valor_base"] = float((list_data[2][39:49]))
                    dic_sts["var1"] = (list_data[2][65:71]).strip()
                    dic_sts["var2"] = (list_data[2][71:77]).strip()
                    dic_sts["var3"] = (list_data[2][77:83]).strip()
                    # Caso node voltage
                    if dic_sts["grandeza"] == 'node  voltage' or dic_sts["grandeza"] == 'branch energy' or dic_sts["grandeza"] == 'branch voltage':
                        dic_sts["var"] = (dic_sts["var1"])[:-1]
                    elif dic_sts["grandeza"] == 'branch current' and dic_sts["var2"] != "":
                        dic_sts["var"] = (dic_sts["var1"])[:-1] + "_" + (dic_sts["var2"])[:-1]
                    elif dic_sts["grandeza"] == 'branch current' and dic_sts["var2"] == "":
                        dic_sts["var"] = (dic_sts["var1"])
                    # elif dic_sts["grandeza"] == 'branch energy':
                    #     dic_sts["var"] = (dic_sts["var1"])[:-1]
                    
                    # Checando se não deu erro
                    is_number = self.check_float(list_data[3][40:54])
                    if is_number:
                        dic_sts["valor_pico"] = float((list_data[3][40:54]))
                        dic_sts["shot"] = (list_data[4][16:20]).strip()
                        dic_sts["var1_shot"] = ((list_data[4][52:99])[:-2]).split("and")[0].replace('"',"").strip()
                        dic_sts["var2_shot"] = ((list_data[4][52:99])[:-2]).split("and")[1].replace('"',"").strip()
                        dic_sts["total_shots"] = (list_data[5][47:52]).strip()
                        if dic_sts["valor_base"] != 0 and dic_sts["valor_base"] != 1:
                            dic_sts["valor_pico_pu"] = dic_sts["valor_pico"] / dic_sts["valor_base"]

                    # Arquiva valores
                    list_data = []
                    list_dics.append(dic_sts)

                # Coleta valores
                if ini_bool:
                    list_data.append(linhas[i])

        df_sts = pd.DataFrame(list_dics)
        df_sts.to_csv(self.path_csv_dfsts, sep=";", decimal=",", encoding="cp1252")  

        return df_sts

    def filter_df_sts(self, df_sts):
        # Curva sempre plotar
        df_temp = df_sts.copy()
        df_temp_1 = df_temp[df_temp["var"].isin(self.curvas_sempre_plotar)]
        # curva_padrao = df_temp["var"]

        # Curva com maior node voltage
        df_temp = df_sts.copy()
        df_temp = df_temp[df_temp["grandeza"] == "node  voltage"]
        df_temp = df_temp[~df_temp["var"].isin(self.curvas_sempre_plotar)]
        df_temp = df_temp.nlargest(self.number_curves_plot, 'valor_pico_pu')
        df_temp_2 = df_temp.dropna(subset=['valor_pico'])
        # curva_node_voltage = df_temp["var"]

        # Curva com maior branch current
        df_temp = df_sts.copy()
        df_temp = df_temp[df_temp["grandeza"] == "branch current"]
        df_temp = df_temp[~df_temp["var"].isin(self.curvas_sempre_plotar)]
        df_temp = df_temp.nlargest(self.number_curves_plot, 'valor_pico_pu')
        df_temp_3 = df_temp.dropna(subset=['valor_pico'])
        # curva_branch_current = df_temp["var1"]

        # Curva com maior branch energy
        df_temp = df_sts.copy()
        df_temp = df_temp[df_temp["grandeza"] == "branch energy"]
        df_temp = df_temp[~df_temp["var"].isin(self.curvas_sempre_plotar)]
        df_temp = df_temp.nlargest(self.number_curves_plot, 'valor_pico_pu')
        df_temp_4 = df_temp.dropna(subset=['valor_pico'])
        # curva_branch_energy = df_temp["var1"]

        # Curva com maior branch voltage
        df_temp = df_sts.copy()
        df_temp = df_temp[df_temp["grandeza"] == "branch voltage"]
        df_temp = df_temp[~df_temp["var"].isin(self.curvas_sempre_plotar)]
        df_temp = df_temp.nlargest(self.number_curves_plot, 'valor_pico_pu')
        df_temp_5 = df_temp.dropna(subset=['valor_pico'])

        union_df = pd.concat([df_temp_1, df_temp_2, df_temp_3, df_temp_4, df_temp_5])
        union_df.to_csv(self.path_csv_dfsts_filt, sep=";", decimal=",", encoding="cp1252")

        return union_df

    def get_dic_sts(self, dic_association, df_sts):
        # Grandezas sempre monitoradas
        df_sts_orig = df_sts.copy()

        # TENSÃO FASE
        df_sts = df_sts_orig.copy()
        df_sts = df_sts[df_sts["grandeza"] == "node  voltage"]
        len_df = len(df_sts[df_sts["var"] == dic_association["v_ft"]])
        if len_df > 0:
            df_sts = df_sts[df_sts["grandeza"] == "node  voltage"]
            v_ft_max = df_sts[df_sts["var"] == dic_association["v_ft"]]["valor_pico"].iloc[0] / 1000
            shot_v_ft_max = df_sts[df_sts["var"] == dic_association["v_ft"]]["shot"].iloc[0]
        else:
            v_ft_max = "-"
            shot_v_ft_max = "-"

        # TENSÃO LINHA
        df_sts = df_sts_orig.copy()
        df_sts = df_sts[df_sts["grandeza"] == "branch voltage"]
        len_df = len(df_sts[df_sts["var"] == dic_association["v_ff"]])
        if len_df > 0:
            df_sts = df_sts[df_sts["grandeza"] == "branch voltage"]
            v_ff_max = df_sts[df_sts["var"] == dic_association["v_ff"]]["valor_pico"].iloc[0] / 1000
            shot_v_ff_max = df_sts[df_sts["var"] == dic_association["v_ff"]]["shot"].iloc[0]
        else:
            v_ff_max = "-"
            shot_v_ff_max = "-"
        
        # CORRENTE RMS
        df_sts = df_sts_orig.copy()
        df_sts = df_sts[df_sts["grandeza"] == "branch current"]
        len_df = len(df_sts[df_sts["var"] == dic_association["cn_rms"]])
        if len_df > 0:
            df_sts = df_sts[df_sts["grandeza"] == "branch current"]
            i_inrush_n_rms = df_sts[df_sts["var"] == dic_association["cn_rms"]]["valor_pico"].iloc[0]
            shot_i_inrush_n_rms = df_sts[df_sts["var"] == dic_association["cn_rms"]]["shot"].iloc[0]
        else:
            i_inrush_n_rms = "-"
            shot_i_inrush_n_rms = "-"

        # CORRENTE FASE MAX
        df_sts = df_sts_orig.copy()
        df_sts = df_sts[df_sts["grandeza"] == "branch current"]
        len_df = len(df_sts[df_sts["var"] == dic_association["cn_max"]])
        if len_df > 0:
            df_sts = df_sts[df_sts["grandeza"] == "branch current"]
            i_inrush_f_max = df_sts[df_sts["var"] == dic_association["cn_max"]]["valor_pico"].iloc[0]
            shot_i_inrush_f_max = df_sts[df_sts["var"] == dic_association["cn_max"]]["shot"].iloc[0]
        else:
            i_inrush_f_max = "-"
            shot_i_inrush_f_max = "-"

        # ENERGIA
        df_sts = df_sts_orig.copy()
        df_sts = df_sts[df_sts["grandeza"] == "branch energy"]
        len_df = len(df_sts[df_sts["var"] == dic_association["energ"]])
        if len_df > 0:
            df_sts = df_sts[df_sts["grandeza"] == "branch energy"]
            energia = df_sts[df_sts["var"] == dic_association["energ"]]["valor_pico"].iloc[0]
            shot_energia = df_sts[df_sts["var"] == dic_association["energ"]]["shot"].iloc[0]
        else:
            energia = "-"
            shot_energia = "-"

        dic_sts_temp = {"equipamento": dic_association["eqp"], "v_ft_max": v_ft_max, "v_ff_max": v_ff_max, "i_inrush_n_rms": i_inrush_n_rms, "i_inrush_f_max": i_inrush_f_max,  "energia": energia}
        dic_shots_temp = {"equipamento": dic_association["eqp"], "v_ft_max": shot_v_ft_max, "v_ff_max": shot_v_ff_max, "i_inrush_n_rms": shot_i_inrush_n_rms, "i_inrush_f_max": shot_i_inrush_f_max,  "energia": shot_energia}

        return dic_sts_temp, dic_shots_temp

    def mount_statisc_df(self, df_sts):
        # Definições iniciais
        list_dic_sts = []
        list_dic_ass = []
        list_dic_shots = []
        used_vars = []
        df_sts_filter = self.filter_df_sts(df_sts)

        # Montando demais trafos observados
        for index, value in df_sts_filter['var'].items():
            if value not in used_vars:
                # Montando dic com valores
                flag_key = False
                if df_sts_filter["grandeza"].loc[index] in ["node  voltage","branch energy","branch voltage"]:
                    key = value
                    flag_key = True

                elif df_sts_filter["grandeza"].loc[index] == "branch current":
                    for chave, dicionario in self.dic_relacoes.items():
                        if "cn_rms" in dicionario and dicionario["cn_rms"] == value and dicionario["v_ft"] in self.dic_relacoes:  # TESTE
                            key = chave
                            flag_key = True
                        elif "cn_max" in dicionario and dicionario["cn_max"] == value:  # TESTE
                            key = chave
                            flag_key = True

                if flag_key == False:
                    print("Atenção! Comportamento inesperado 001")

                dic_sts_temp, dic_shots_temp = self.get_dic_sts(self.dic_relacoes[key], df_sts)
                # Salvando resultados
                list_dic_sts.append(dic_sts_temp)     
                list_dic_ass.append(self.dic_relacoes[key])           
                list_dic_shots.append(dic_shots_temp)

                # Definição das variáveis
                var_v_ft_max = self.dic_relacoes[key]["v_ft"]
                var_v_ff_max = self.dic_relacoes[key]["v_ff"]
                var_i_inrush_n_rms = self.dic_relacoes[key]["cn_rms"]
                var_i_inrush_f_max = self.dic_relacoes[key]["cn_max"]
                var_energia = self.dic_relacoes[key]["energ"]
                used_vars.append(var_v_ft_max)
                used_vars.append(var_v_ff_max)
                used_vars.append(var_i_inrush_n_rms)
                used_vars.append(var_i_inrush_f_max)
                used_vars.append(var_energia)

        # Montar tabela
        df_var = pd.DataFrame(list_dic_ass)
        df_resultados = pd.DataFrame(list_dic_sts)
        df_shot = pd.DataFrame(list_dic_shots)
        df_var.columns = ['Equipamento','Tensão Fase-Terra Máxima (kV)','Tensão Fase-Fase Máxima (kV)','Corrente Inrush Neutro RMS (A)','Corrente Inrush Fase Máxima (A)','Energia (kJ)']
        df_resultados.columns = ['Equipamento','Tensão Fase-Terra Máxima (kV)','Tensão Fase-Fase Máxima (kV)','Corrente Inrush Neutro RMS (A)','Corrente Inrush Fase Máxima (A)','Energia (kJ)']
        df_shot.columns = ['Equipamento','Tensão Fase-Terra Máxima (kV)','Tensão Fase-Fase Máxima (kV)','Corrente Inrush Neutro RMS (A)','Corrente Inrush Fase Máxima (A)','Energia (kJ)']

        df_var.to_csv(self.path_csv_var, sep=";", decimal=",", encoding="cp1252")
        df_resultados.to_csv(self.path_csv_res, sep=";", decimal=",", encoding="cp1252")
        df_shot.to_csv(self.path_csv_sht, sep=";", decimal=",", encoding="cp1252")

        return df_var, df_var, df_shot

    #==============================================================================================
    #
    # FUNC 03
    #
    #==============================================================================================
    def run_ATP_shot(self, list_shots = []):
        """
        Realiza a execução dos shots resultantes do estatístico
        """
        # Busca df_shots na pasta
        if len(list_shots) == 0:
            if os.path.exists(self.path_csv_sht):
                df_shots = pd.read_csv(self.path_csv_sht, sep=";", decimal=",", encoding="cp1252")
            else:
                print("Arquivo CSV dos Shots não localizado")

            # Pegando a lista de shots a serem rodados
            df_num_shots = df_shots.iloc[:, -5:]
            list_shots = []
            for coluna in df_num_shots.columns:
                list_shots.extend(df_num_shots[coluna].unique())
            list_shots = [str(valor) for valor in list_shots]
            list_shots = [valor for valor in list_shots if valor != "-"]
            list_shots = list(set(list_shots))
            list_shots.sort()
        else:
            pass

        # if "-" not in self.path_fileATP:
        for i, shot in enumerate(list_shots):
            print(f"...Executando shot {i+1}/{len(list_shots)}")
            shot_name = int(shot)
            shot_name = f"{self.folder_atp}shot{shot_name:04}.dat"
            output_file_shot = self.output_file[:-4] + f"_shot{shot}.dat"

            self.adjust_time_simulations("shot", shot_name)
            self.run_ATP(shot_name, output_file_shot, make_backup=False)

        return

    #==============================================================================================
    #
    # FUNC 04
    #
    #==============================================================================================
    def converte_ATP_shot(self):
        """
        Realiza a execução dos shots resultantes do estatístico
        """
        # Busca pl4 na pasta
        file_list = os.listdir(self.folder_atp)
        pl4_files = [file for file in file_list if file.upper().endswith(".PL4")]
        dic_MAT = {}
        for i, shot in enumerate(pl4_files):
            if shot[:4] == "shot":
                if "-" not in self.folder_atp + f"/{shot}" and "." not in self.folder_atp:
                    shot_pl4 = self.folder_atp + f"{shot}"
                    shot_mat = self.folder_atp + f"{shot[:-4]}.mat"
                    comando = f'{self.path_mat} "{shot_pl4}"'

                    p = subprocess.Popen(comando)
                    p.wait()
                    p.terminate()                     
                else:
                    print("ERRO NA CONVERSÃO! Foi encontrado o caracter '-' ou '.' no diretório de trabalho. Elimine ou substitua-o para prosseguir.")
        return


    #==============================================================================================
    #
    # FUNC 05
    #
    #==============================================================================================
    def carrega_ATP_shot_MAT(self):
        """
        Realiza a execução dos shots resultantes do estatístico
        """
        # Busca pl4 na pasta
        file_list = os.listdir(self.folder_atp)
        mat_files = [file for file in file_list if file.upper().endswith(".MAT")]
        dic_MAT = {}
        for i, shot in enumerate(mat_files):
            res_Atp = sio.loadmat(self.folder_atp + shot)
            res_Atp = {chave.upper(): valor for chave, valor in res_Atp.items()}

            shot_num = str(int(shot[4:8]))
            dic_MAT[shot_num] = res_Atp
        return dic_MAT
    
    def mount_trifasic_df(self, t, fase_a, fase_b, fase_c, title, tipo, nfase_a = "Fase A", nfase_b = "Fase B", nfase_c = "Fase C"):
        combined_array = np.concatenate((t, fase_a, fase_b, fase_c), axis=1)
        df = pd.DataFrame(combined_array, columns=pd.MultiIndex.from_tuples([
            (tipo, title, "t"),
            (tipo, title, nfase_a),
            (tipo, title, nfase_b),
            (tipo, title, nfase_c),]))
        df.set_index((tipo, title, "t"), inplace=True)
        df.index.name = None
    
        return df

    def mount_monofasic_df(self, t, fase_a, title, tipo):
        combined_array = np.concatenate((t, fase_a), axis=1)
        df = pd.DataFrame(combined_array, columns=pd.MultiIndex.from_tuples([
            (tipo, title, "t"),
            (tipo, title, title)]))
        df.set_index((tipo, title, "t"), inplace=True)
        df.index.name = None
        
        return df

    def get_df_shots(self, save_df = True):
        """
        Realiza a execução dos shots resultantes do estatístico
        """
        dic_MAT = self.carrega_ATP_shot_MAT()

        # Verifica se dfs estão na pasta
        if os.path.exists(self.path_csv_sht):
            df_shots = pd.read_csv(self.path_csv_sht, sep=";", decimal=",", encoding="cp1252")
            df_shots = df_shots.iloc[:, 2:]
        else:
            print("Arquivo CSV dos Shots não localizado")
            return            
        if os.path.exists(self.path_csv_var):
            df_vars = pd.read_csv(self.path_csv_var, sep=";", decimal=",", encoding="cp1252")
            df_vars = df_vars.iloc[:, 2:]
        else:
            print("Arquivo CSV das variáveis dos Shots não localizado")
            return

        # Começa a iterar sobre o dataframe de variáveis e ir montando nosso dataframe com os plots
        df = pd.DataFrame()
        for index, row in df_vars.iterrows():
            for tipo, grandeza in row.items():
                if grandeza == grandeza:
                    if "-" in grandeza:
                        grandeza = grandeza.replace("-","_")
                    if tipo == 'Tensão Fase-Terra Máxima (kV)':
                        # Verifica se é TF_D para pegar o shot completo
                        if len(self.dic_relacao_regime) > 0 and grandeza in  self.dic_relacao_regime:
                            # Identificação dos Arrays
                            fase_a = "V" + self.dic_relacao_regime[grandeza] + "A"
                            fase_b = "V" + self.dic_relacao_regime[grandeza] + "B"
                            fase_c = "V" + self.dic_relacao_regime[grandeza] + "C"
                        # elif "-" in grandeza:
                        #     grandeza = grandeza.replace("-","_")
                        #     fase_a = "V" + grandeza + "A"
                        #     fase_b = "V" + grandeza + "B"
                        #     fase_c = "V" + grandeza + "C"                        
                        else:
                            # Identificação dos Arrays
                            fase_a = "V" + grandeza + "A"
                            fase_b = "V" + grandeza + "B"
                            fase_c = "V" + grandeza + "C"

                        shot = str(df_shots[tipo].loc[index])
                        # Montando dataframe
                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "node  voltage")
                        df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Tensão Fase-Fase Máxima (kV)':
                        # Verifica se é TF_D para pegar o shot completo
                        if len(self.dic_relacao_regime) > 0 and grandeza in self.dic_relacao_regime:
                            # Identificação dos Arrays
                            fase_a = "V" + self.dic_relacao_regime[grandeza] + "A"
                            fase_b = "V" + self.dic_relacao_regime[grandeza] + "B"
                            fase_c = "V" + self.dic_relacao_regime[grandeza] + "C"
                        else:
                            # Identificação dos Arrays
                            fase_a = f"V{grandeza}A{grandeza}C"
                            fase_b = f"V{grandeza}B{grandeza}A"
                            fase_c = f"V{grandeza}C{grandeza}B"

                        shot = str(df_shots[tipo].loc[index])
                        # Montando dataframe
                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "branch voltage", nfase_a="Fase AC", nfase_b="Fase BA", nfase_c="Fase CB")
                        df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Corrente Inrush Neutro RMS (A)':
                        # Identificação dos Arrays
                        fase_at = "I" + grandeza + "TERRA"
                        fase_ta = "ITERRA" + grandeza

                        shot = str(df_shots[tipo].loc[index])
                        if shot != "-":
                            fase_a = ""
                            if fase_at in dic_MAT[shot].keys():
                                fase_a = fase_at
                            elif fase_ta in dic_MAT[shot].keys():
                                fase_a = fase_ta

                            if fase_a != "":
                                df_temp = self.mount_monofasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], grandeza, "branch current N")
                                df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Corrente Inrush Fase Máxima (A)':
                        # Identificação dos Arrays
                        nos = grandeza.split("_")

                        fase_a = f"I{nos[0]}A{nos[1]}A"
                        fase_b = f"I{nos[0]}B{nos[1]}B"
                        fase_c = f"I{nos[0]}C{nos[1]}C"
                        shot = str(df_shots[tipo].loc[index])

                        # Montando dataframe
                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "branch current F")
                        df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Energia (kJ)':
                        # Identificação dos Arrays
                        # Identificação dos Arrays
                        fase_at = "E" + grandeza + "ATERRA"
                        fase_ta = "ETERRA" + grandeza + "A"

                        if fase_at in dic_MAT[shot].keys():
                            fase_a = fase_at
                            fase_b = "E" + grandeza + "BTERRA"
                            fase_c = "E" + grandeza + "CTERRA"
                        elif fase_ta in dic_MAT[shot].keys():
                            fase_a = fase_ta
                            fase_b = "ETERRA" + grandeza + "B"
                            fase_c = "ETERRA" + grandeza + "C"
                        else:
                            pass
                        shot = str(df_shots[tipo].loc[index])

                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "energy")
                        df = pd.concat([df, df_temp], axis=1)

                    else:
                        pass
        # Escrevendo os resultados
        if save_df:
            df.to_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=True, index=True)
        # dfcsv = pd.read_csv(self.path_csv_plt, index_col=[0], header=[0,1,2], sep=";", decimal=",", encoding="cp1252")
        return df
    
    def separate_list_3f(self, plots_var):
        resultados = []
        for i in range(len(plots_var) - 1):
            for j in range(i + 1, len(plots_var)):
                if plots_var[i][:-1] == plots_var[j][:-1]:
                    resultados.append(plots_var[i][:-1])
                    break
        resultados_unicos = list(set(resultados))
        return resultados_unicos, resultados_unicos
    
    def get_df_shots_total(self, df_sts):
        """
        Realiza a execução dos shots resultantes do estatístico
        """
        dic_MAT = self.carrega_ATP_shot_MAT()
        df = pd.DataFrame()

        # Verifica se dfs estão na pasta
        for shot in dic_MAT:
            df_sts_shot = df_sts[df_sts["shot"] == shot]

            # Varrendo o dataframe do shot específico
            for index, row in df_sts_shot.iterrows():
                # Caso seja tensão
                if row["grandeza"] == "node  voltage":
                    # Se grandeza trifásica
                    if row["var3"] != "":
                        grandeza = row["var"].replace("-","_")
                        fase_a = f"V{grandeza}A"
                        fase_b = f"V{grandeza}B"
                        fase_c = f"V{grandeza}C"

                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "node  voltage")
                        df = pd.concat([df, df_temp], axis=1)                        
                    else:
                        pass

                elif row["grandeza"] == "branch voltage":
                    # Se grandeza trifásica
                    grandeza = row["var"].replace("-","_")
                    fase_a = f"V{grandeza}A{grandeza}C"
                    fase_b = f"V{grandeza}B{grandeza}A"
                    fase_c = f"V{grandeza}C{grandeza}B"

                    df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "branch voltage", nfase_a = "Fase AC", nfase_b = "Fase BA", nfase_c = "Fase CB")
                    df = pd.concat([df, df_temp], axis=1)                        

                # Caso seja energia
                elif row["grandeza"] == "branch energy":
                    grandeza = row["var"].replace("-","_")
                    fase_at = "E" + grandeza + "ATERRA"
                    fase_ta = "ETERRA" + grandeza + "A"
                    if fase_at in dic_MAT[shot].keys():
                        fase_a = fase_at
                        fase_b = "E" + grandeza + "BTERRA"
                        fase_c = "E" + grandeza + "CTERRA"
                    elif fase_ta in dic_MAT[shot].keys():
                        fase_a = fase_ta
                        fase_b = "ETERRA" + grandeza + "B"
                        fase_c = "ETERRA" + grandeza + "C"
                    else:
                        pass
                    df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "energy")
                    df = pd.concat([df, df_temp], axis=1)
                elif row["grandeza"] == "branch current":
                    # Se corrente de fase
                    if row["var2_shot"] != "":
                        grandeza = row['var'].replace("-","_")
                        grandeza_1 = row["var1_shot"][:-1]
                        grandeza_2 = row["var2_shot"][:-1]
                        fase_a = f"I{grandeza_1}A{grandeza_2}A"
                        fase_b = f"I{grandeza_1}B{grandeza_2}B"
                        fase_c = f"I{grandeza_1}C{grandeza_2}C"

                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "branch current F")
                        df = pd.concat([df, df_temp], axis=1)                        
                    else:
                        # Se corrente de neutro
                        grandeza = row["var"].replace("-","_")
                        fase_at = "I" + grandeza + "TERRA"
                        fase_ta = "ITERRA" + grandeza
                        fase_a = ""
                        if fase_at in dic_MAT[shot].keys():
                            fase_a = fase_at
                        elif fase_ta in dic_MAT[shot].keys():
                            fase_a = fase_ta

                        if fase_a != "":
                            df_temp = self.mount_monofasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], grandeza, "branch current N")
                            df = pd.concat([df, df_temp], axis=1)
                else:
                    pass
                    print(f"Row index: {index}")

        # Salva dataframe final
        # df = df.rename_axis("Tempo")
        df = df.sort_index()
        df = df.fillna(method='ffill')
        df.to_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=True, index=True)
        return df



        if os.path.exists(self.path_csv_sht):
            df_shots = pd.read_csv(self.path_csv_sht, sep=";", decimal=",", encoding="cp1252")
            df_shots = df_shots.iloc[:, 2:]
        else:
            print("Arquivo CSV dos Shots não localizado")
            return            
        if os.path.exists(self.path_csv_var):
            df_vars = pd.read_csv(self.path_csv_var, sep=";", decimal=",", encoding="cp1252")
            df_vars = df_vars.iloc[:, 2:]
        else:
            print("Arquivo CSV das variáveis dos Shots não localizado")
            return

        # Começa a iterar sobre o dataframe de variáveis e ir montando nosso dataframe com os plots
        df = pd.DataFrame()
        for index, row in df_vars.iterrows():
            for tipo, grandeza in row.items():
                if grandeza == grandeza:
                    if "-" in grandeza:
                        grandeza = grandeza.replace("-","_")
                    if tipo == 'Tensão Fase-Terra Máxima (kV)':
                        # Verifica se é TF_D para pegar o shot completo
                        if len(self.dic_relacao_regime) > 0 and grandeza in  self.dic_relacao_regime:
                            # Identificação dos Arrays
                            fase_a = "V" + self.dic_relacao_regime[grandeza] + "A"
                            fase_b = "V" + self.dic_relacao_regime[grandeza] + "B"
                            fase_c = "V" + self.dic_relacao_regime[grandeza] + "C"
                        # elif "-" in grandeza:
                        #     grandeza = grandeza.replace("-","_")
                        #     fase_a = "V" + grandeza + "A"
                        #     fase_b = "V" + grandeza + "B"
                        #     fase_c = "V" + grandeza + "C"                        
                        else:
                            # Identificação dos Arrays
                            fase_a = "V" + grandeza + "A"
                            fase_b = "V" + grandeza + "B"
                            fase_c = "V" + grandeza + "C"

                        shot = str(df_shots[tipo].loc[index])
                        # Montando dataframe
                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "node  voltage")
                        df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Tensão Fase-Fase Máxima (kV)':
                        # Verifica se é TF_D para pegar o shot completo
                        if len(self.dic_relacao_regime) > 0 and grandeza in self.dic_relacao_regime:
                            # Identificação dos Arrays
                            fase_a = "V" + self.dic_relacao_regime[grandeza] + "A"
                            fase_b = "V" + self.dic_relacao_regime[grandeza] + "B"
                            fase_c = "V" + self.dic_relacao_regime[grandeza] + "C"
                        else:
                            # Identificação dos Arrays
                            fase_a = f"V{grandeza}A{grandeza}C"
                            fase_b = f"V{grandeza}B{grandeza}A"
                            fase_c = f"V{grandeza}C{grandeza}B"

                        shot = str(df_shots[tipo].loc[index])
                        # Montando dataframe
                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "branch voltage", nfase_a="Fase AC", nfase_b="Fase BA", nfase_c="Fase CB")
                        df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Corrente Inrush Neutro RMS (A)':
                        # Identificação dos Arrays
                        fase_at = "I" + grandeza + "TERRA"
                        fase_ta = "ITERRA" + grandeza

                        shot = str(df_shots[tipo].loc[index])
                        if shot != "-":
                            fase_a = ""
                            if fase_at in dic_MAT[shot].keys():
                                fase_a = fase_at
                            elif fase_ta in dic_MAT[shot].keys():
                                fase_a = fase_ta

                            if fase_a != "":
                                df_temp = self.mount_monofasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], grandeza, "branch current N")
                                df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Corrente Inrush Fase Máxima (A)':
                        # Identificação dos Arrays
                        nos = grandeza.split("_")

                        fase_a = f"I{nos[0]}A{nos[1]}A"
                        fase_b = f"I{nos[0]}B{nos[1]}B"
                        fase_c = f"I{nos[0]}C{nos[1]}C"
                        shot = str(df_shots[tipo].loc[index])

                        # Montando dataframe
                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "branch current F")
                        df = pd.concat([df, df_temp], axis=1)

                    elif tipo == 'Energia (kJ)':
                        # Identificação dos Arrays
                        # Identificação dos Arrays
                        fase_at = "E" + grandeza + "ATERRA"
                        fase_ta = "ETERRA" + grandeza + "A"

                        if fase_at in dic_MAT[shot].keys():
                            fase_a = fase_at
                            fase_b = "E" + grandeza + "BTERRA"
                            fase_c = "E" + grandeza + "CTERRA"
                        elif fase_ta in dic_MAT[shot].keys():
                            fase_a = fase_ta
                            fase_b = "ETERRA" + grandeza + "B"
                            fase_c = "ETERRA" + grandeza + "C"
                        else:
                            pass
                        shot = str(df_shots[tipo].loc[index])

                        df_temp = self.mount_trifasic_df(dic_MAT[shot]["T"], dic_MAT[shot][fase_a], dic_MAT[shot][fase_b], dic_MAT[shot][fase_c], grandeza, "energy")
                        df = pd.concat([df, df_temp], axis=1)

                    else:
                        pass
        # Escrevendo os resultados
        if save_df:
            df.to_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=True, index=True)
        # dfcsv = pd.read_csv(self.path_csv_plt, index_col=[0], header=[0,1,2], sep=";", decimal=",", encoding="cp1252")
        return df

    #==============================================================================================
    #
    # FUNC 06
    #
    #==============================================================================================
    def plot_df_line_chart(self, df, x_label, y_label, title, path_save="", apply_envoltoria=False, fator_env=1, show = False):
        # Define se a envoltória vai ser impressa para TFs
        df = df.rename_axis(None) #.sort_index()
        if apply_envoltoria:                
            envoltoria_sup = np.empty(df.index.shape, dtype=np.float32)
            envoltoria_sup[df.index <= 0.1667] = 2.0 * fator_env                              # 10 ciclos 
            envoltoria_sup[(df.index > 0.1667) & (df.index <= 0.3333)] = 1.82 * fator_env     # 20 ciclos
            envoltoria_sup[(df.index > 0.3333) & (df.index <= 1.6667)] = 1.50 * fator_env     # 100 ciclos
            envoltoria_sup[(df.index > 1.6667) & (df.index <= 3.6000)] = 1.40 * fator_env
            envoltoria_sup[(df.index > 3.6)    & (df.index <= 10)    ] = 1.35 * fator_env
            envoltoria_sup[(df.index > 10)     & (df.index <= 20)    ] = 1.25 * fator_env
            envoltoria_sup[(df.index > 20)     & (df.index <= 60)    ] = 1.20 * fator_env
            envoltoria_sup[(df.index > 60)     & (df.index <= 480)   ] = 1.15 * fator_env
            envoltoria_sup[df.index > 480] = 1.1 * fator_env
            envoltoria_inf = - envoltoria_sup

        # Iniciando a plotagem
        df.plot(kind='line')
        if apply_envoltoria:
            plt.plot(df.index,envoltoria_sup, label = "", color="black")
            plt.plot(df.index,envoltoria_inf, label = "", color="black")

        plt.xlim(df.index.min(), df.index.max())
        plt.xlabel(x_label)
        plt.ylabel(y_label)
        if title != "":
            plt.title(title)
        plt.legend(loc='upper right')
        plt.tight_layout(pad=0.0, w_pad=0.0, h_pad=0.0)

        # Salvando o png
        if path_save != "":
            if os.path.isfile(path_save): 
                os.remove(path_save)
            plt.gcf().set_size_inches(6, 3.54)
            plt.tight_layout()
            plt.savefig(path_save, dpi=self.dpi_res)
        
        if show:
            plt.show(block=False)
            plt.legend(list(df.columns.get_level_values(0)), loc='upper right')

        return        
        
    def generate_plots(self, save_plot=True, kv_pu=False, apply_envoltoria=True):
        if os.path.exists(self.path_csv_plt):
            df_plots = pd.read_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=[0,1,2])
            df_plots = df_plots.set_index(df_plots.columns[0]).rename_axis("Tempo")
        else:
            print("Arquivo CSV com dados de plotagem não localizado!")
            return
        # df_plots = self.get_df_shots(save_df=False)

        # Verifica se dfs estão na pasta
        if os.path.exists(self.path_csv_dfsts):
            df_sts = pd.read_csv(self.path_csv_dfsts, sep=";", decimal=",", encoding="cp1252")
            df_sts = df_sts.iloc[:, 1:]
        else:
            print("Arquivo CSV dos Shots não localizado")
            return            
        if os.path.exists(self.path_csv_var):
            df_vars = pd.read_csv(self.path_csv_var, sep=";", decimal=",", encoding="cp1252")
            df_vars = df_vars.iloc[:, 1:]
        else:
            print("Arquivo CSV das variáveis dos Shots não localizado")
            return       
        
        folder_pics = self.folder_atp + "pngs\\"
        if not os.path.exists(folder_pics):
            os.makedirs(folder_pics)

        # Começa a montar os plots
        counter = 0
        for index, row in df_vars.iterrows():
            for tipo, grandeza in row.items():
                if grandeza == grandeza:
                    grandeza = grandeza.replace("-","_")
                    if tipo == 'Tensão Fase-Terra Máxima (kV)':
                        # Define o dataframe de dados
                        df_temp = df_sts[(df_sts["grandeza"] == "node  voltage") & (df_sts["var"] == grandeza)]
                        if len(df_temp) > 0:
                            vbase = (df_temp["valor_base"].iloc[0])
                            if kv_pu:                           
                                fator_env = 1
                                df_plot = df_plots["node  voltage"][grandeza] / vbase
                                title = f'Tensão Nó (p.u.): {row["Equipamento"]}'
                                path_savefig = folder_pics + f'tensao-{counter}-{grandeza}-pu.png'
                                y_axis = 'Tensão (p.u.)'
                            else:
                                fator_env = (df_temp["valor_base"].iloc[0]) / 1000
                                df_plot = df_plots["node  voltage"][grandeza] / 1000
                                title = f'Tensão Nó (kV): {row["Equipamento"]}'
                                path_savefig = folder_pics + f'tensao-{counter}-{grandeza}-kV.png'
                                y_axis = 'Tensão (kV)'

                            # Imprime o dataframe de dados                        
                            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=True, fator_env=fator_env)
                            counter += 1 

                            # Visualizando a forma inicial                        
                            last_2_percent = int(len(df_plot) * self.slice_df_visualiza)
                            filtered_df = df_plot.head(last_2_percent)
                            title = f'Tensão Nó (kV): {row["Equipamento"]} (Inicial)'
                            if kv_pu:                           
                                path_savefig = folder_pics + f'tensao-{counter}-{grandeza}-pu_inicial.png'
                            else:
                                path_savefig = folder_pics + f'tensao-{counter}-{grandeza}-kV_inicial.png'
                            self.plot_df_line_chart(filtered_df, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=True, fator_env=fator_env)
                            counter += 1 

                            # Visualizando a forma em regime
                            last_2_percent = int(len(df_plot) * self.slice_df_visualiza)
                            filtered_df = df_plot.tail(last_2_percent)
                            title = f'Tensão Nó (kV): {row["Equipamento"]} (Regime)'
                            if kv_pu:                           
                                path_savefig = folder_pics + f'tensao-{counter}-{grandeza}-pu_regime.png'
                            else:
                                path_savefig = folder_pics + f'tensao-{counter}-{grandeza}-kV_regime.png'
                            self.plot_df_line_chart(filtered_df, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=True, fator_env=fator_env)
                            counter += 1
                    
                    elif tipo == 'Corrente Inrush Neutro RMS (A)':
                        # Define o dataframe de dados
                        df_temp = df_sts[(df_sts["grandeza"] == "branch current") & (df_sts["var"] == grandeza)]
                        if len(df_temp) > 0:
                            vbase = (df_temp["valor_base"].iloc[0])

                            # Definindo plot
                            df_plot = df_plots["branch current N"][grandeza]
                            title = f'Corrente no Neutro em RMS (A): {row["Equipamento"]}'
                            path_savefig = folder_pics + f'corrente_neutro-{counter}-{grandeza}-A.png'
                            y_axis = 'Corrente RMS (A)'

                            # Imprime o dataframe de dados                        
                            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False,)
                            counter += 1 

                    elif tipo == 'Corrente Inrush Fase Máxima (A)':
                        # Define o dataframe de dados
                        df_temp = df_sts[(df_sts["grandeza"] == "branch current") & (df_sts["var"] == grandeza)]
                        if len(df_temp) > 0:
                            vbase = (df_temp["valor_base"].iloc[0])

                            # Definindo plot
                            df_plot = df_plots["branch current F"][grandeza]
                            title = f'Corrente de Fase (A): {row["Equipamento"]}'
                            path_savefig = folder_pics + f'corrente_fase-{counter}-{grandeza}-A.png'
                            y_axis = 'Corrente (A)'

                            # Imprime o dataframe de dados                        
                            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False,)
                            counter += 1 

                            # Visualizando a forma inicial
                            last_2_percent = int(len(df_plot) * self.slice_df_visualiza)
                            filtered_df = df_plot.head(last_2_percent)
                            title = f'Corrente de Fase (A): {row["Equipamento"]} (Inicial)'
                            path_savefig = folder_pics + f'corrente_fase-{counter}-{grandeza}-A_inicial.png'
                            self.plot_df_line_chart(filtered_df, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False)
                            counter += 1 

                            # Visualizando a forma em regime
                            last_2_percent = int(len(df_plot) * self.slice_df_visualiza)
                            filtered_df = df_plot.tail(last_2_percent)
                            title = f'Corrente de Fase (A): {row["Equipamento"]} (Regime)'
                            path_savefig = folder_pics + f'corrente_fase-{counter}-{grandeza}-A_regime.png'
                            self.plot_df_line_chart(filtered_df, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False)
                            counter += 1 

                    elif tipo == 'Energia (MJ)':
                        # Define o dataframe de dados
                        df_temp = df_sts[(df_sts["grandeza"] == "branch energy") & (df_sts["var"] == grandeza)]
                        if len(df_temp) > 0:
                            vbase = (df_temp["valor_base"].iloc[0])

                            df_plot = df_plots["energy"][grandeza] / vbase
                            title = f'Energia (MJ): {row["Equipamento"]}'
                            path_savefig = folder_pics + f'energia-{counter}-{grandeza}-MJ.png'
                            y_axis = 'Energia (MJ)'

                            # Imprime o dataframe de dados                        
                            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False)

                            counter += 1 
                    
                    else:
                        pass

        # Gerar Curva com todos I neutro
        df_plot = df_plots["branch current N"]
        title = f'Corrente no Neutro em RMS (A): {row["Equipamento"]}'
        path_savefig = folder_pics + f'corrente_neutro-{counter}-TOTAL.png'
        y_axis = 'Corrente RMS (A)'
        self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False,)
        return
    
    def generate_plots_all(self, save_plot=True, kv_pu=False, apply_envoltoria=True, df_plots=pd.DataFrame(), df_sts=pd.DataFrame()):       
        folder_pics = self.folder_atp + "pngs\\"
        if not os.path.exists(folder_pics):
            os.makedirs(folder_pics)

        # Começa a montar os plots            
        level_0_headers = list(df_plots.columns.get_level_values(0).unique())
        level_1_headers = list(df_plots.columns.get_level_values(1).unique())
        counter = 0

        # Começa o filtro
        for index, grandeza in enumerate(level_0_headers):
            df_plots_grandeza = df_plots[grandeza].drop_duplicates()
            level_1_headers = list(df_plots_grandeza.columns.get_level_values(0).unique())

            if grandeza == "node  voltage":
                for subindex, var in enumerate(level_1_headers): 
                    df_sts_temp = df_sts[(df_sts["grandeza"] == grandeza) & (df_sts["var"].str.replace("-","_") == var)]
                    vbase = (df_sts_temp["valor_base"].iloc[0])
                    # if kv_pu:                           
                    fator_env = 1
                    df_plot = df_plots_grandeza[var] / vbase
                    title = f'Tensão Nó (p.u.): {var}'
                    path_savefig = folder_pics + f'tensao-{counter}-{var}-pu.png'
                    y_axis = 'Tensão (p.u.)'
                    self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=True, fator_env=fator_env, show=True)
                    counter += 1
                    # else:
                    fator_env = vbase / 1000
                    df_plot = df_plots_grandeza[var] / 1000
                    title = f'Tensão Nó (kV): {var}'
                    path_savefig = folder_pics + f'tensao-{counter}-{grandeza}-kV.png'
                    y_axis = 'Tensão (kV)'                       
                    self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=True, fator_env=fator_env, show=True)
                    counter += 1

            elif grandeza == "energy":
                for subindex, var in enumerate(level_1_headers): 
                    df_sts_temp = df_sts[(df_sts["grandeza"] == "branch energy") & (df_sts["var"].str.replace("-","_") == var)]
                    vbase = (df_sts_temp["valor_base"].iloc[0])
                    #
                    df_plot = df_plots_grandeza[var] / 1000
                    title = f'Energia (kJ): {var}'
                    path_savefig = folder_pics + f'energia-{counter}-{grandeza}-kJ.png'
                    y_axis = 'Energia (kJ)'                    
                    self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)
                    counter += 1 

            elif grandeza == "branch current F":
                for subindex, var in enumerate(level_1_headers): 
                    df_sts_temp = df_sts[(df_sts["grandeza"] == "branch current") & (df_sts["var"].str.replace("-","_") == var)]
                    vbase = (df_sts_temp["valor_base"].iloc[0])
                    #
                    df_plot = df_plots_grandeza[var]
                    title = f'Corrente de Fase (A): {var}'
                    path_savefig = folder_pics + f'corrente_fase-{counter}-{grandeza}-A.png'
                    y_axis = 'Corrente (A)'                    
                    self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)
                    counter += 1 

            elif grandeza == "branch current N":
                for subindex, var in enumerate(level_1_headers): 
                    df_sts_temp = df_sts[(df_sts["grandeza"] == "branch current") & (df_sts["var"].str.replace("-","_") == var)]
                    vbase = (df_sts_temp["valor_base"].iloc[0])
                    #
                    df_plot = df_plots_grandeza[var]
                    title = f'Corrente no Neutro em RMS (A): {var}'
                    path_savefig = folder_pics + f'corrente_neutro-{counter}-{grandeza}-A.png'
                    y_axis = 'Corrente RMS (A)'                    
                    self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)
                    counter += 1 
                    
            else:
                pass
      
        # Gerar Curva com todos I neutro
        if "branch current N" in level_0_headers:
            df_plot = df_plots["branch current N"]
            title = f'Corrente no Neutro em RMS (A)'
            path_savefig = folder_pics + f'corrente_neutro-{counter}-TOTAL.png'
            y_axis = 'Corrente RMS (A)'
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)
        return

    def generate_plot_specific(self, save_plot=True, kv_pu=False, apply_envoltoria=True, saida_ATP_custom="", grandeza = "", var = "", drop_plot=[]):  
        #
        # Possibilita o programa a ler arquivo de saída gerado fora dele
        if saida_ATP_custom != "":
            self.output_file = saida_ATP_custom
        folder_pics = self.folder_atp + "pngs\\"
        if not os.path.exists(folder_pics):
            os.makedirs(folder_pics)

        df_sts = self.get_df_sts()
        df_plots = pd.read_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=[0,1,2])
        df_plots = df_plots.set_index(df_plots.columns[0]).rename_axis("Tempo")
     
        # Começa a montar os plots            
        level_0_headers = list(df_plots.columns.get_level_values(0).unique())
        level_1_headers = list(df_plots.columns.get_level_values(1).unique())

        if grandeza not in level_0_headers:
            print("Valor setado para grandeza é inválido!")
            return
        if var not in level_1_headers and grandeza != "branch current N":
            print("Valor setado para var é inválido!")
            return

        # Começa o filtro
        df_plots_grandeza = df_plots[grandeza].drop_duplicates()

        if grandeza == "node  voltage":
            df_sts_temp = df_sts[(df_sts["grandeza"] == grandeza) & (df_sts["var"].str.replace("-","_") == var)]
            vbase = (df_sts_temp["valor_base"].iloc[0])
            # if kv_pu:                           
            fator_env = 1
            df_plot = df_plots_grandeza[var] / vbase
            title = ""#f'Tensão Nó (p.u.): {var}'
            path_savefig = folder_pics + f'tensao-{var}-pu.png'
            y_axis = 'Tensão (p.u.)'
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=apply_envoltoria, fator_env=fator_env, show=True)
            # else:
            fator_env = vbase / 1000
            df_plot = df_plots_grandeza[var] / 1000
            title = ""#f'Tensão Nó (kV): {var}'
            path_savefig = folder_pics + f'tensao-{grandeza}-kV.png'
            y_axis = 'Tensão (kV)'                       
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=apply_envoltoria, fator_env=fator_env, show=True)

        elif grandeza == "branch voltage":
            df_sts_temp = df_sts[(df_sts["grandeza"] == grandeza) & (df_sts["var"].str.replace("-","_") == var)]
            vbase = (df_sts_temp["valor_base"].iloc[0])
            # if kv_pu:                           
            fator_env = 1
            df_plot = df_plots_grandeza[var] / vbase
            title = ""#f'Tensão Nó (p.u.): {var}'
            path_savefig = folder_pics + f'tensao-{var}-pu.png'
            y_axis = 'Tensão (p.u.)'
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=apply_envoltoria, fator_env=fator_env, show=True)
            # else:
            fator_env = vbase / 1000
            df_plot = df_plots_grandeza[var] / 1000
            title = ""#f'Tensão Nó (kV): {var}'
            path_savefig = folder_pics + f'tensao-{grandeza}-kV.png'
            y_axis = 'Tensão (kV)'                       
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=apply_envoltoria, fator_env=fator_env, show=True)

        elif grandeza == "energy":
            df_sts_temp = df_sts[(df_sts["grandeza"] == "branch energy") & (df_sts["var"].str.replace("-","_") == var)]
            vbase = (df_sts_temp["valor_base"].iloc[0])
            #
            df_plot = df_plots_grandeza[var] / 1000
            title = ""#f'Energia (kJ): {var}'
            path_savefig = folder_pics + f'energia-{grandeza}-kJ.png'
            y_axis = 'Energia (kJ)'                    
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)

        elif grandeza == "branch current F":           
            df_sts_temp = df_sts[(df_sts["grandeza"] == "branch current") & (df_sts["var"].str.replace("-","_") == var)]
            vbase = (df_sts_temp["valor_base"].iloc[0])
            #
            df_plot = df_plots_grandeza[var]
            title = ""#f'Corrente de Fase (A): {var}'
            path_savefig = folder_pics + f'corrente_fase-{grandeza}-A.png'
            y_axis = 'Corrente (A)'                    
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)

        elif grandeza == "branch current N":
            if var != "":
                df_sts_temp = df_sts[(df_sts["grandeza"] == "branch current") & (df_sts["var"].str.replace("-","_") == var)]
                vbase = (df_sts_temp["valor_base"].iloc[0])
                df_plot = df_plots_grandeza[var]
            else:
                df_sts_temp = df_sts[(df_sts["grandeza"] == "branch current") & (~df_sts["var"].isin(drop_plot))]
                columns_df = list(df_plots_grandeza.columns.get_level_values(1).unique())
                columns_df = [element for element in columns_df if element not in drop_plot]

                df_plot = df_plots_grandeza[columns_df]
                max_values = df_plot.max()
                sorted_columns = max_values.sort_values(ascending=False).index
                df_plot = df_plot[sorted_columns]
            
            title = ""#f'Corrente no Neutro em RMS (A): {var}'
            path_savefig = folder_pics + f'corrente_neutro-{grandeza}-A.png'
            y_axis = 'Corrente RMS (A)'                    
            self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)
                
        else:
            pass
      
        # Gerar Curva com todos I neutro
        # df_plot = df_plots["branch current N"]
        # title = f'Corrente no Neutro em RMS (A)'
        # path_savefig = folder_pics + f'corrente_neutro-{counter}-TOTAL.png'
        # y_axis = 'Corrente RMS (A)'
        # self.plot_df_line_chart(df_plot, 'Tempo (s)', y_axis, title, path_savefig, apply_envoltoria=False, show=True)
        return


    #==============================================================================================
    #
    # FUNC 07
    #
    #==============================================================================================
    def generate_report_tem(self):
        # Carrega plots com base nos dados da pasta
        if os.path.exists(self.path_csv_plt):
            df_plots = pd.read_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=[0,1,2])
            df_plots = df_plots.set_index(df_plots.columns[0]).rename_axis("Tempo")
            # df_plots = df_plots.iloc[:, 1:]
        else:
            print("Arquivo CSV com dados de plotagem não localizado!")
            return
        # df_plots = self.get_df_shots(save_df=False)
        # df_sts = self.get_df_sts()
        if os.path.exists(self.path_csv_dfsts):
            df_sts = pd.read_csv(self.path_csv_dfsts, sep=";", decimal=",", encoding="cp1252")
            df_sts = df_sts.iloc[:, 1:]
        else:
            print("Arquivo CSV dos resultados estatísticos não localizado")
            return
        if os.path.exists(self.path_csv_var):
            df_vars = pd.read_csv(self.path_csv_var, sep=";", decimal=",", encoding="cp1252")
            df_vars = df_vars.iloc[:, 1:]
        else:
            print("Arquivo CSV das variáveis dos Shots não localizado")
            return
        if os.path.exists(self.path_csv_res):
            df_resultados = pd.read_csv(self.path_csv_res, sep=";", decimal=",", encoding="cp1252")
            df_resultados = df_resultados.iloc[:, 1:]
        else:
            print("Arquivo CSV das variáveis dos Shots não localizado")
            return

        # Inicia execução do Relatório
        document = Document(self.path_word_template_report)
        counter_fig = 1

        # Tamanho figura
        altura_fig_inch = Inches(3.54)
        comprim_fig_inch = Inches(6)

        # Cria bloco de título
        paragraph = document.add_paragraph(style="tmp_Estilo_1")
        paragraph.add_run("RELATÓRIO AUTOMÁTICO – SIMULAÇÃO ATP VIA PYTHON")
        paragraph = document.add_paragraph(style="tmp_Estilo_2")
        paragraph.add_run("Resultados estatísticos e determinísticos")
        equipamento = df_vars["Equipamento"].iloc[0]
        paragraph = document.add_paragraph(style="tmp_Estilo_3")
        paragraph.add_run(f"Energização do equipamento – {equipamento}")
        
        # Item Cenários
        paragraph = document.add_paragraph(style="tmp_Estilo_4")
        paragraph.add_run("Cenários")            
        paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
        paragraph.add_run(f"Figura 1-{counter_fig}: Cenário 1")
        paragraph = document.add_paragraph(style="onsFigura")
        paragraph.add_run("")
        paragraph = document.add_paragraph(style="onsIlustração_Fim")
        paragraph.add_run("")
        counter_fig += 1

        # Item Resultados Estatísticos
        paragraph = document.add_paragraph(style="tmp_Estilo_4")
        paragraph.add_run("Resultados estatísticos")
        paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
        paragraph.add_run("Tabela 1-1: Resultados estatísticos do cenário 1 – Tensão pré-manobra: [X,XXX] p.u.")

        ## Adiciona tabela de resultados estatísticos
        table = document.add_table(rows=1, cols=len(df_resultados.columns))
        table.autofit = False
        # table.style = 'Table Grid'
        for i, col_name in enumerate(df_resultados.columns):
            table.cell(0, i).text = col_name
        header_row = table.rows[0]
        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)

        # Ajustando df_resultados para valores numéricos
        df_resultados = df_resultados.replace('-', np.nan)
        for col in df_resultados.columns[1:]:
            df_resultados[col] = df_resultados[col].astype(float)

        for _, row in df_resultados.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                if isinstance(value, float):
                    value = round(value,1)
                    if value != value:
                        value = "-"
                row_cells[i].text = str(value).replace(".", ",")
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row_cells[i].horizontal_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = document.add_paragraph(style="onsIlustração_Fim")
        paragraph.add_run("Notas: 1. Tensões base: Fase-terra XXX kVpico, XXX kVrms e fase-fase XXX kVpico.")

        # Add a page break
        page_break = document.add_paragraph()
        run = page_break.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Item Resultados Determinísticos
        paragraph = document.add_paragraph(style="tmp_Estilo_4")
        paragraph.add_run("Resultados determinísticos")

        folder_pics = self.folder_atp + "pngs\\"
        if not os.path.exists(folder_pics):
            os.makedirs(folder_pics)

        # Cria bloco por equipamento
        count_pg_break = 0
        for index, row in df_vars.iterrows():
            equipamento = row["Equipamento"]
            tensao_fase = row['Tensão Fase-Terra Máxima (kV)']
            tensao_linha = row['Tensão Fase-Fase Máxima (kV)']
            corrente_neutro = row['Corrente Inrush Neutro RMS (A)']
            corrente_fase = row['Corrente Inrush Fase Máxima (A)']
            energia = row['Energia (kJ)']
            # Verifica possibilidade de -
            tensao_fase_mod = tensao_fase.replace("-", "_") if "-" in str(tensao_fase) else tensao_fase
            tensao_linha_mod = tensao_linha.replace("-", "_") if "-" in str(tensao_linha) else tensao_linha
            corrente_neutro_mod = corrente_neutro.replace("-", "_") if "-" in str(corrente_neutro) else corrente_neutro
            corrente_fase_mod = corrente_fase.replace("-", "_") if "-" in str(corrente_fase) else corrente_fase
            energia_mod = energia.replace("-", "_") if "-" in str(energia) else energia
            fator_env = df_sts[(df_sts["grandeza"] == "node  voltage") & (df_sts["var"] == tensao_fase)]["valor_base"].iloc[0]/1000

            # Cria título do item
            paragraph = document.add_paragraph(style="tmp_Estilo_5")
            paragraph.add_run(equipamento)

            #======================================================================================
            # Monta a figura referente a máxima tensão instantânea
            df_temp = df_plots["node  voltage"][tensao_fase_mod]/1000
            # title = f'Tensão Nó (kV): {tensao_fase}A - {tensao_fase}B - {tensao_fase}C'
            # title = f'Tensão Nó (kV): {equipamento}'
            path_savefig = folder_pics + 'fig_temp.png'
            self.plot_df_line_chart(df_temp, 'Tempo (s)', "Tensão (kV)", "", path_savefig, apply_envoltoria=True, fator_env=fator_env)

            # Insere figura no documento
            paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
            paragraph.add_run(f"Figura 1-{counter_fig}: Máxima tensão fase instantânea - {equipamento}")
            document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
            paragraph = document.add_paragraph(style="onsIlustração_Fim")
            paragraph.add_run(" ")
            counter_fig += 1
            count_pg_break += 1

            # # Add a page break
            if count_pg_break == 2:
                last_paragraph = document.paragraphs[-1]
                last_run = last_paragraph.runs[-1]
                last_run.add_break(WD_BREAK.PAGE)
                # page_break = document.add_paragraph()
                # run = page_break.add_run()
                # run.add_break(WD_BREAK.PAGE)
                count_pg_break = 0

            # Monta a figura referente a máxima tensão instantânea - porrada inicial
            last_2_percent = int(len(df_temp) * self.slice_df_visualiza)
            filtered_df = df_temp.head(last_2_percent)
            # title = f'Tensão Nó (kV): {tensao_fase}A - {tensao_fase}B - {tensao_fase}C (Instante Inicial)'
            title = f'Tensão Nó (kV): {equipamento} (Instante Inicial)'
            self.plot_df_line_chart(filtered_df, 'Tempo (s)', "Tensão (kV)", "", path_savefig, apply_envoltoria=True, fator_env=fator_env)

            # Insere figura no documento
            paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
            paragraph.add_run(f"Figura 1-{counter_fig}: Máxima tensão fase instantânea - Detalhe da energização - {equipamento}")
            document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
            paragraph = document.add_paragraph(style="onsIlustração_Fim")
            paragraph.add_run(" ")
            counter_fig += 1
            count_pg_break += 1

            # # Add a page break
            if count_pg_break == 2:                
                last_paragraph = document.paragraphs[-1]
                last_run = last_paragraph.runs[-1]
                last_run.add_break(WD_BREAK.PAGE)
                # page_break = document.add_paragraph()
                # run = page_break.add_run()
                # run.add_break(WD_BREAK.PAGE)
                count_pg_break = 0

            # Monta a figura referente a máxima tensão instantânea - regime
            last_2_percent = int(len(df_temp) * self.slice_df_visualiza)
            filtered_df = df_temp.tail(last_2_percent)
            # title = f'Tensão Nó (kV): {tensao_fase}A - {tensao_fase}B - {tensao_fase}C (Regime)'
            title = f'Tensão Nó (kV): {equipamento} (Regime)'
            self.plot_df_line_chart(filtered_df, 'Tempo (s)', "Tensão (kV)", "", path_savefig, apply_envoltoria=True, fator_env=fator_env)

            # Insere figura no documento
            paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
            paragraph.add_run(f"Figura 1-{counter_fig}: Máxima tensão fase instantânea - Detalhe do regime - {equipamento}")
            document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
            paragraph = document.add_paragraph(style="onsIlustração_Fim")
            paragraph.add_run(" ")
            counter_fig += 1
            count_pg_break += 1

            # # Add a page break
            if count_pg_break == 2:
                last_paragraph = document.paragraphs[-1]
                last_run = last_paragraph.runs[-1]
                last_run.add_break(WD_BREAK.PAGE)
                # page_break = document.add_paragraph()
                # run = page_break.add_run()
                # run.add_break(WD_BREAK.PAGE)
                count_pg_break = 0

            #======================================================================================
            # Monta a figura referente a máxima tensão instantânea
            if "branch voltage" in df_plots.columns and tensao_fase_mod in df_plots["branch voltage"].columns:
                df_temp = df_plots["branch voltage"][tensao_fase_mod]/1000
                # title = f'Tensão Nó (kV): {tensao_fase}A - {tensao_fase}B - {tensao_fase}C'
                # title = f'Tensão Nó (kV): {equipamento}'
                path_savefig = folder_pics + 'fig_temp.png'
                self.plot_df_line_chart(df_temp, 'Tempo (s)', "Tensão (kV)", "", path_savefig, apply_envoltoria=True, fator_env=fator_env)

                # Insere figura no documento
                paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                paragraph.add_run(f"Figura 1-{counter_fig}: Máxima tensão linha instantânea - {equipamento}")
                document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                paragraph = document.add_paragraph(style="onsIlustração_Fim")
                paragraph.add_run(" ")
                counter_fig += 1
                count_pg_break += 1

                # # Add a page break
                if count_pg_break == 2:
                    last_paragraph = document.paragraphs[-1]
                    last_run = last_paragraph.runs[-1]
                    last_run.add_break(WD_BREAK.PAGE)
                    # page_break = document.add_paragraph()
                    # run = page_break.add_run()
                    # run.add_break(WD_BREAK.PAGE)
                    count_pg_break = 0

                # Monta a figura referente a máxima tensão instantânea - porrada inicial
                last_2_percent = int(len(df_temp) * self.slice_df_visualiza)
                filtered_df = df_temp.head(last_2_percent)
                # title = f'Tensão Nó (kV): {tensao_fase}A - {tensao_fase}B - {tensao_fase}C (Instante Inicial)'
                title = f'Tensão Fase (kV): {equipamento} (Instante Inicial)'
                self.plot_df_line_chart(filtered_df, 'Tempo (s)', "Tensão (kV)", "", path_savefig, apply_envoltoria=True, fator_env=fator_env)

                # Insere figura no documento
                paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                paragraph.add_run(f"Figura 1-{counter_fig}: Máxima tensão linha instantânea - Detalhe da energização - {equipamento}")
                document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                paragraph = document.add_paragraph(style="onsIlustração_Fim")
                paragraph.add_run(" ")
                counter_fig += 1
                count_pg_break += 1

                # # Add a page break
                if count_pg_break == 2:                
                    last_paragraph = document.paragraphs[-1]
                    last_run = last_paragraph.runs[-1]
                    last_run.add_break(WD_BREAK.PAGE)
                    # page_break = document.add_paragraph()
                    # run = page_break.add_run()
                    # run.add_break(WD_BREAK.PAGE)
                    count_pg_break = 0

                # Monta a figura referente a máxima tensão instantânea - regime
                last_2_percent = int(len(df_temp) * self.slice_df_visualiza)
                filtered_df = df_temp.tail(last_2_percent)
                # title = f'Tensão Nó (kV): {tensao_fase}A - {tensao_fase}B - {tensao_fase}C (Regime)'
                title = f'Tensão Fase (kV): {equipamento} (Regime)'
                self.plot_df_line_chart(filtered_df, 'Tempo (s)', "Tensão (kV)", "", path_savefig, apply_envoltoria=True, fator_env=fator_env)

                # Insere figura no documento
                paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                paragraph.add_run(f"Figura 1-{counter_fig}: Máxima tensão linha instantânea - Detalhe do regime - {equipamento}")
                document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                paragraph = document.add_paragraph(style="onsIlustração_Fim")
                paragraph.add_run(" ")
                counter_fig += 1
                count_pg_break += 1

                # # Add a page break
                if count_pg_break == 2:
                    last_paragraph = document.paragraphs[-1]
                    last_run = last_paragraph.runs[-1]
                    last_run.add_break(WD_BREAK.PAGE)
                    # page_break = document.add_paragraph()
                    # run = page_break.add_run()
                    # run.add_break(WD_BREAK.PAGE)
                    count_pg_break = 0

            #======================================================================================
            # Monta a figura referente a máxima corrente de fase
            if "branch current F" in df_plots.columns:
                if corrente_fase_mod in df_plots["branch current F"]:
                    df_temp = df_plots["branch current F"][corrente_fase_mod]
                    # title = f'Corrente de Fase (A): {corrente_fase}'
                    title = f'Corrente de Fase (A): {equipamento}'
                    path_savefig = folder_pics + 'fig_temp.png'
                    self.plot_df_line_chart(df_temp, 'Tempo (s)', "Corrente Pico (A)", "", path_savefig, apply_envoltoria=False)

                    # Insere figura no documento
                    paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                    paragraph.add_run(f"Figura 1-{counter_fig}: Máxima corrente de fase - {equipamento}")
                    document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                    paragraph = document.add_paragraph(style="onsIlustração_Fim")
                    paragraph.add_run(" ")
                    counter_fig += 1
                    count_pg_break += 1

                    # # Add a page break
                    if count_pg_break == 2:
                        last_paragraph = document.paragraphs[-1]
                        last_run = last_paragraph.runs[-1]
                        last_run.add_break(WD_BREAK.PAGE)
                        # page_break = document.add_paragraph()
                        # run = page_break.add_run()
                        # run.add_break(WD_BREAK.PAGE)
                        count_pg_break = 0

                    # Monta a figura referente a máxima corrente de fase - porrada inicial
                    last_2_percent = int(len(df_temp) * self.slice_df_visualiza)
                    filtered_df = df_temp.head(last_2_percent)
                    # title = f'Corrente de Fase (A): {corrente_fase} (Instante Inicial)'
                    title = f'Corrente de Fase (A): {equipamento} (Instante Inicial)'
                    self.plot_df_line_chart(filtered_df, 'Tempo (s)', "Corrente Pico (A)", "", path_savefig, apply_envoltoria=False)

                    # Insere figura no documento
                    paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                    paragraph.add_run(f"Figura 1-{counter_fig}: Máxima corrente de fase - Detalhe da energização - {equipamento}")
                    document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                    paragraph = document.add_paragraph(style="onsIlustração_Fim")
                    paragraph.add_run(" ")
                    counter_fig += 1
                    count_pg_break += 1

                    # # Add a page break
                    if count_pg_break == 2:
                        last_paragraph = document.paragraphs[-1]
                        last_run = last_paragraph.runs[-1]
                        last_run.add_break(WD_BREAK.PAGE)
                        # page_break = document.add_paragraph()
                        # run = page_break.add_run()
                        # run.add_break(WD_BREAK.PAGE)
                        count_pg_break = 0

                    # Monta a figura referente a máxima corrente de fase - regime
                    last_2_percent = int(len(df_temp) * self.slice_df_visualiza)
                    filtered_df = df_temp.tail(last_2_percent)
                    # title = f'Corrente de Fase (A): {corrente_fase} (Regime)'
                    title = f'Corrente de Fase (A): {equipamento} (Regime)'
                    self.plot_df_line_chart(filtered_df, 'Tempo (s)', "Corrente Pico (A)", "", path_savefig, apply_envoltoria=False)

                    # Insere figura no documento
                    paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                    paragraph.add_run(f"Figura 1-{counter_fig}: Máxima corrente de fase - Regime - {equipamento}")
                    document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                    paragraph = document.add_paragraph(style="onsIlustração_Fim")
                    paragraph.add_run(" ")
                    counter_fig += 1
                    count_pg_break += 1

                    # # Add a page break
                    if count_pg_break == 2:
                        last_paragraph = document.paragraphs[-1]
                        last_run = last_paragraph.runs[-1]
                        last_run.add_break(WD_BREAK.PAGE)
                        # page_break = document.add_paragraph()
                        # run = page_break.add_run()
                        # run.add_break(WD_BREAK.PAGE)
                        count_pg_break = 0

            #======================================================================================
            # Monta a figura referente a máxima corrente de fase
            if "branch current N" in df_plots.columns:
                if corrente_neutro_mod in df_plots["branch current N"]:
                    df_temp = df_plots["branch current N"][corrente_neutro_mod]
                    # title = f'Corrente de Neutro (A): {corrente_neutro}'
                    title = f'Corrente de Neutro (A): {equipamento}'
                    path_savefig = folder_pics + 'fig_temp.png'
                    self.plot_df_line_chart(df_temp, 'Tempo (s)', "Corrente Neutro RMS (A)", "", path_savefig, apply_envoltoria=False)

                    # Insere figura no documento
                    paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                    paragraph.add_run(f"Figura 1-{counter_fig}: Máxima corrente de neutro - {equipamento}")
                    document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                    paragraph = document.add_paragraph(style="onsIlustração_Fim")
                    paragraph.add_run(" ")
                    counter_fig += 1
                    count_pg_break += 1

                    # # Add a page break
                    if count_pg_break == 2:
                        last_paragraph = document.paragraphs[-1]
                        last_run = last_paragraph.runs[-1]
                        last_run.add_break(WD_BREAK.PAGE)
                        # page_break = document.add_paragraph()
                        # run = page_break.add_run()
                        # run.add_break(WD_BREAK.PAGE)
                        count_pg_break = 0

            # Monta a figura referente a máxima energia
            if "energy" in df_plots.columns:
                if energia_mod in df_plots["energy"]:
                    df_temp = df_plots["energy"][energia_mod]/1000
                    # title = f'Energia (MJ): {energia}'
                    title = f'Energia (kJ): {equipamento}'
                    path_savefig = folder_pics + 'fig_temp.png'
                    self.plot_df_line_chart(df_temp, 'Tempo (s)', "Energia (kJ)", "", path_savefig, apply_envoltoria=False)

                    # Insere figura no documento
                    paragraph = document.add_paragraph(style="onsIlustração_Tabela_Início")
                    paragraph.add_run(f"Figura 1-{counter_fig}: Máxima energia absorvida pelo para-raios - {equipamento}")
                    document.add_picture(path_savefig, width=comprim_fig_inch, height=altura_fig_inch)
                    paragraph = document.add_paragraph(style="onsIlustração_Fim")
                    paragraph.add_run(" ")
                    counter_fig += 1
                    count_pg_break += 1

                    # # Add a page break
                    if count_pg_break == 2:
                        last_paragraph = document.paragraphs[-1]
                        last_run = last_paragraph.runs[-1]
                        last_run.add_break(WD_BREAK.PAGE)
                        # page_break = document.add_paragraph()
                        # run = page_break.add_run()
                        # run.add_break(WD_BREAK.PAGE)
                        count_pg_break = 0

            # # Add a page break
            if count_pg_break != 0:
                last_paragraph = document.paragraphs[-1]
                last_run = last_paragraph.runs[-1]
                last_run.add_break(WD_BREAK.PAGE)
                # page_break = document.add_paragraph()
                # run = page_break.add_run()
                # run.add_break(WD_BREAK.PAGE)
                count_pg_break = 0

        paragraph = document.add_paragraph(style="tmp_Estilo_4")
        paragraph.add_run("Comentários e conclusões")

        conclusion = ""
        conclusion += f"Na energização dos transformadores, as sobretensões máximas de manobra são inferiores ao nível básico de impulso de "
        conclusion += f"manobra típico dos equipamentos de 230 kV envolvidos, tais como transformadores, equipamentos terminais e equipamentos" 
        conclusion += f"de barra, e inferior ao fator de surto típico utilizado em linhas de 230 kV."
        paragraph = document.add_paragraph(style="onsNormal")
        paragraph.add_run(conclusion)
        conclusion = ""
        conclusion += f"As correntes de inrush máximas observadas na energização são inferiores à suportabilidade mecânica típica dos transformadores" 
        conclusion += f"neste porte. As correntes de neutro apresentam um bom amortecimento, não sendo esperada atuação de proteção de sobrecorrente de neutro."
        conclusion += f"A energia absorvida nos para-raios é consideravelmente inferior à sua capacidade."
        paragraph = document.add_paragraph(style="onsNormal")
        paragraph.add_run(conclusion)
        conclusion = ""
        conclusion += f"Portanto, não são observadas violações de suportabilidade de equipamentos para as manobras de energização de transformadores nas "
        conclusion += f"condições e topologias previstas nesta área de recomposição, as quais se mostram viáveis nos critérios preconizados nos Procedimentos de Rede e na ótica da engenharia."
        paragraph = document.add_paragraph(style="onsNormal")
        paragraph.add_run(conclusion)

        # Save the Word document
        document.save(self.path_docx_report)
        if self.open_word:
            try:
                subprocess.run(["start", self.path_docx_report], shell=True)
            except FileNotFoundError:
                print("Error: Could not open the file. Make sure it exists and you have the required application (e.g., Microsoft Word) installed.")
        pass

    #==============================================================================================
    #
    # FUNC 01_07
    #
    #==============================================================================================
    def remove_files_plot(self):
        for filename in os.listdir(self.folder_atp):
            if filename.upper().endswith(".PL4") or filename.upper().endswith(".MAT"):
                file_path = os.path.join(self.folder_atp, filename)
                try:
                    os.remove(file_path)
                except:
                    pass

    def roda_estudo_ATP(self,   rodar_ATP = True, 
                                tempo_ATP = 1, 
                                num_simul = 2, 
                                saida_ATP_custom = "", 
                                rodar_SHOT = True, 
                                tempo_SHOT = 1, 
                                converter_SHOT = True,
                                gerar_plots = True,
                                montar_csv_resultados = False,
                                remover_plots = True,
                                gera_relatorio = True):
        ## Função para rodar o caso ATP informado
        time_start = time.time()
        time_func_inic = time.time()
        decimals = 3
        print("Execução do estudo estatístico do ATP via Python iniciado...")
        print(f"Arquivo analisado: {self.path_fileATP}")
        counter_func = 1
        if rodar_ATP:
            print(f"Função {counter_func}: Modificando o número de simulações para {num_simul}.")
            counter_func += 1
            # Ajustar Número & Tempo das Simulações            
            self.adjust_num_simulations(num_simul)
            print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            time_func_inic = time.time()

            print(f"Função {counter_func}: Modificando o tempo de simulação para {self.time_simulation}.")
            counter_func += 1
            self.time_simulation = tempo_ATP
            self.adjust_time_simulations(type_file="main")
            print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            time_func_inic = time.time()

            # Executa a simulação
            print(f"Função {counter_func}: Executando o caso ATP.")
            counter_func += 1
            self.run_ATP()
            print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            time_func_inic = time.time()

        # Possibilita o programa a ler arquivo de saída gerado fora dele
        if saida_ATP_custom != "":
            self.output_file = saida_ATP_custom

        ## Função para coletar resultados estatísticos
        print(f"Função {counter_func}: Coletando os resultados estatísticos")
        counter_func += 1
        df_sts = self.get_df_sts()

        # Traz os resultados estatísticos para uma forma mais amigável 

        if len(df_sts) > 0:       
            df_var, df_resultados, df_shots = self.mount_statisc_df(df_sts)
            print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            time_func_inic = time.time()

            ## Função para rodar os shots
            if rodar_SHOT:
                print(f"Função {counter_func}: Iniciando a execução dos SHOTS selecionados")
                counter_func += 1
                self.time_simulation_shot = tempo_SHOT
                self.run_ATP_shot()
                converter_SHOT = True
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            # Função para converter arquivos .PL4 para .MAT
            if converter_SHOT:
                print(f"Função {counter_func}: Iniciando a conversão dos shots de .PL4 para .MAT")
                counter_func += 1
                # Converte os shots na pasta
                self.converte_ATP_shot()
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            # Montar os dataframes com dados de plotagem 
            # Converte os shots na pasta
            print(f"Função {counter_func}: Montar dataframe com os dados de plotagem")
            counter_func += 1

            if rodar_ATP or rodar_SHOT or converter_SHOT or montar_csv_resultados:
                df_plots = self.get_df_shots()
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            elif os.path.exists(self.path_csv_plt):
                df_plots = pd.read_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=[0,1,2])
                df_plots = df_plots.set_index(df_plots.columns[0]).rename_axis("Tempo")
                # df_plots = df_plots.iloc[:, 1:]
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            else:
                print("Arquivo CSV dos Resultados Simulação não localizado")
                return
            
            time_func_inic = time.time()
            # Função para montar os plots
            # self.slice_df_visualiza = 0.1
            if gerar_plots:
                print(f"Função {counter_func}: Gerar visualizações na pasta")
                counter_func += 1
                self.generate_plots(save_plot=True, kv_pu=True, apply_envoltoria=True)
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            # Função para montar o relatório 
            if gera_relatorio:
                print(f"Função {counter_func}: Gerando o relatório de estudo")
                counter_func += 1
                self.generate_report_tem()
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            if remover_plots:
                print(f"Função {counter_func}: Limpando arquivos .PL4 e .MAT gerados")
                counter_func += 1
                self.remove_files_plot()
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            print(f"Programa executado com sucesso em {round(time.time()-time_start,decimals)} segundos!")
        
        else:
            print(f"Programa não executado com sucesso! Verificar arquivo de saída, pois o estatístico falhou!")

    def roda_estudo_ATP_full(self,  rodar_ATP = True, 
                                    tempo_ATP = 1, 
                                    num_simul = 2, 
                                    saida_ATP_custom = "", 
                                    rodar_SHOT = True, 
                                    tempo_SHOT = 1, 
                                    converter_SHOT = True,
                                    gerar_plots = True,
                                    montar_csv_resultados = False,
                                    remover_plots = True,
                                    gera_relatorio = True):
        ## Função para rodar o caso ATP informado
        time_start = time.time()
        time_func_inic = time.time()
        decimals = 3
        print("Execução do estudo estatístico do ATP via Python iniciado...")
        print(f"Arquivo analisado: {self.path_fileATP}")
        counter_func = 1
        if rodar_ATP:
            print(f"Função {counter_func}: Modificando o número de simulações para {num_simul}.")
            counter_func += 1
            # Ajustar Número & Tempo das Simulações            
            self.adjust_num_simulations(num_simul)
            print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            time_func_inic = time.time()

            print(f"Função {counter_func}: Modificando o tempo de simulação para {self.time_simulation}.")
            counter_func += 1
            self.time_simulation = tempo_ATP
            self.adjust_time_simulations(type_file="main")
            print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            time_func_inic = time.time()

            # Executa a simulação
            print(f"Função {counter_func}: Executando o caso ATP.")
            counter_func += 1
            self.run_ATP()
            print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
            time_func_inic = time.time()

        # Possibilita o programa a ler arquivo de saída gerado fora dele
        if saida_ATP_custom != "":
            self.output_file = saida_ATP_custom

        ## Função para coletar resultados estatísticos
        print(f"Função {counter_func}: Coletando os resultados estatísticos")
        counter_func += 1
        df_sts = self.get_df_sts()        
        
        if len(df_sts) > 0:
            df_sts = df_sts.dropna(subset=["shot"])
            ## Rodar TODOS os shots
            if rodar_SHOT:
                print(f"Função {counter_func}: Iniciando a execução dos SHOTS selecionados")
                counter_func += 1
                self.time_simulation_shot = tempo_SHOT
                list_shots = list(set(list(df_sts["shot"])))
                self.run_ATP_shot(list_shots)
                converter_SHOT = True
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            # Função para converter arquivos .PL4 para .MAT
            if converter_SHOT:
                print(f"Função {counter_func}: Iniciando a conversão dos shots de .PL4 para .MAT")
                counter_func += 1
                # Converte os shots na pasta
                self.converte_ATP_shot()
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            if montar_csv_resultados:
                print(f"Função {counter_func}: Iniciando a conversão dos shots de .MAT para .CSV")
                counter_func += 1
                #
                df_plots = self.get_df_shots_total(df_sts)
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()
            else:
                df_plots = pd.read_csv(self.path_csv_plt, sep=";", decimal=",", encoding="cp1252", header=[0,1,2])
                df_plots = df_plots.set_index(df_plots.columns[0]).rename_axis("Tempo")
                # df_plots = df_plots.iloc[:, 1:]
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")

            if gerar_plots:
                print(f"Função {counter_func}: Gerar visualizações na pasta")
                counter_func += 1
                self.generate_plots_all(save_plot=True, kv_pu=True, apply_envoltoria=True, df_plots=df_plots, df_sts=df_sts)
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()
            

            # Função para montar o relatório 
            if gera_relatorio:
                print(f"Função {counter_func}: Gerando o relatório de estudo")
                counter_func += 1
                self.generate_report_tem()
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            if remover_plots:
                print(f"Função {counter_func}: Limpando arquivos .PL4 e .MAT gerados")
                counter_func += 1
                self.remove_files_plot()
                print(f"...função executada em {round(time.time()-time_func_inic,decimals)} segundos!")
                time_func_inic = time.time()

            print(f"Programa executado com sucesso em {round(time.time()-time_start,decimals)} segundos!")
            
        else:
            print(f"Programa não executado com sucesso! Verificar arquivo de saída, pois o estatístico falhou!")



if __name__ == "__main__":
    #==============================================================================
    # Sumário de funções
    func_01 = 0             # Função para rodar o ATP
    func_02 = 0             # Função para coletar resultados estatísticos
    func_03 = 0             # Função para rodar os shots
    func_04 = 0             # Função para converter arquivos .PL4 para .MAT
    func_05 = 0             # Função carregar os dados .MAT para memória
    func_06 = 0             # Função para montar os plots
    func_07 = 0             # Função para montar o relatório
    func_01_07 = 1          # Função para rodar tudo!
    #==============================================================================
    # Arquivo para simulação
    path_fileATP = r"D:\Sync\2_resultados\1ºTFA_CCO_230.atp"
    oATP = ATP(path_fileATP)

    # Parâmetros do estudo
    oATP.dic_relacao_regime = {"TF_D": "TF"}
    oATP.dic_relacoes = {
               "TF":{"eqp": "TF1 230/138kV Gaspar 2", "v_ft": "TF", "cn_rms": "NTF_R", "cn_max": "TF_TC", "energ": "PR"},
               "TF_D":{"eqp": "TF1 230/138kV Gaspar 2 (Regime)", "v_ft": "TF_D", "cn_rms": "", "cn_max": "", "energ": ""},
            "CBA52":{"eqp": "TF1 525/230kV Curitiba", "v_ft": "CBA52", "cn_rms": "NCBA_R", "cn_max": "", "energ": ""},
            "ARE52":{"eqp": "TF1 525/230kV Areia", "v_ft": "ARE52", "cn_rms": "NARE_R", "cn_max": "", "energ": ""},
            "UMB23":{"eqp": "TF1 230/69kV Umbará", "v_ft": "UMB23", "cn_rms": "NUMB_R", "cn_max": "", "energ": ""},
            "UBE23":{"eqp": "TF1 230/69kV Uberaba", "v_ft": "UBE23", "cn_rms": "NUBE_R", "cn_max": "", "energ": ""},
            "CBC23":{"eqp": "TF1 230/69kV Curitiba Centro", "v_ft": "CBC23", "cn_rms": "NCBC_R", "cn_max": "", "energ": ""},
            "BAT52":{"eqp": "TF1 525/230kV Bateias", "v_ft": "BAT52", "cn_rms": "NBAT_R", "cn_max": "", "energ": ""},
            "BAT23":{"eqp": "TF1 230/138kV Bateias", "v_ft": "BAT23", "cn_rms": "NBA2_R", "cn_max": "", "energ": ""},
              "816":{"eqp": "TF1 230/69kV Campo Comprido", "v_ft": "816", "cn_rms": "NCCO_R", "cn_max": "", "energ": ""},
            "CIC23":{"eqp": "TF1 230/69kV Cidade Industrial Curitiba", "v_ft": "CIC23", "cn_rms": "NCIC_R", "cn_max": "", "energ": ""},
             "9542":{"eqp": "TF1 525/230kV Curitiba Leste", "v_ft": "9542", "cn_rms": "NCTL_R", "cn_max": "", "energ": ""},
              "829":{"eqp": "TF1 230/138kV Ponta Grossa Norte", "v_ft": "829", "cn_rms": "NPGN_R", "cn_max": "", "energ": ""},
              "834":{"eqp": "TF1 230/34,5kV São Mateus do Sul", "v_ft": "834", "cn_rms": "NSMS_R", "cn_max": "", "energ": ""},
             "1015":{"eqp": "TF1 230/138kV Joinville", "v_ft": "1015", "cn_rms": "NJOI_R", "cn_max": "", "energ": ""},
              "939":{"eqp": "TF1 230/138kV Blumenau", "v_ft": "939", "cn_rms": "NBL2_R", "cn_max": "", "energ": ""},
              "938":{"eqp": "TF1 525/230kV Blumenau", "v_ft": "938", "cn_rms": "NBL5_R", "cn_max": "", "energ": ""},
              "991":{"eqp": "TF1 230/138kV Itajaí 2", "v_ft": "991", "cn_rms": "NITJ_R", "cn_max": "", "energ": ""},
             "9533":{"eqp": "TF1 230/138kV Gaspar 2", "v_ft": "9533", "cn_rms": "NGAS_R", "cn_max": "", "energ": ""},
            }
    
    # Inicializando as funções
    #-------------------------------------------------------------------------------
    ## Funções porcionadas
    ## Função para rodar o caso ATP informado
    if func_01 == 1:       
        # Ajustar Número & Tempo das Simulações
        oATP.adjust_num_simulations(num_simul = 3)
        oATP.adjust_time_simulations(type_file="main")

        # Executa a simulação
        oATP.run_ATP()

    ## Função para coletar resultados estatísticos
    if func_02 == 1:
        # Leitura para dataframe dos resultados estatísticos
        df_sts = oATP.get_df_sts()

        # Traz os resultados estatísticos para uma forma mais amigável        
        df_var, df_resultados, df_shots = oATP.mount_statisc_df(df_sts)

    ## Função para rodar os shots
    if func_03 == 1:
        # Executar os shots na pasta
        oATP.time_simulation_shot = 1
        oATP.run_ATP_shot()

    # Função para converter arquivos .PL4 para .MAT
    if func_04 == 1:
        # Converte os shots na pasta
        oATP.converte_ATP_shot()

    # Montar os dataframes com dados de plotagem 
    if func_05 == 1:
        # Converte os shots na pasta
        df_plots = oATP.get_df_shots()

    # Função para montar os plots
    if func_06 == 1:
        # oATP.slice_df_visualiza = 0.1
        oATP.generate_plots(save_plot=True, kv_pu=True, apply_envoltoria=True)

    # Função para montar o relatório 
    if func_07 == 1:
        # oATP.slice_df_visualiza = 0.1
        oATP.generate_report_tem()

    # Função para montar o relatório 
    if func_07 == 1:
        # oATP.slice_df_visualiza = 0.1
        oATP.generate_report_tem()

    #-------------------------------------------------------------------------------
    ## Funções agrupadas
    if func_01_07:
        # Função para executar o estudo ATP de maneira completa
        oATP.roda_estudo_ATP()
    pass
"""Tests for the run module."""

import docx
import pytest

from cmi_docx import run, styles


def test_find_run_lt_same_paragraph() -> None:
    """Test that comparing FindRun of the same paragraph works."""
    document = docx.Document()
    paragraph = document.add_paragraph("Hello, world!")

    find_run1 = run.FindRun(paragraph, (0, 1), (0, 5))
    find_run2 = run.FindRun(paragraph, (1, 2), (5, 10))

    assert find_run1 < find_run2


def test_find_run_lt_different_paragraphs() -> None:
    """Test that comparing FindRun of different paragraphs fails."""
    document = docx.Document()
    paragraph1 = document.add_paragraph("Hello, world!")
    paragraph2 = document.add_paragraph("Hello, world!")

    find_run1 = run.FindRun(paragraph1, (0, 1), (0, 5))
    find_run2 = run.FindRun(paragraph2, (0, 1), (0, 5))

    with pytest.raises(ValueError):
        assert find_run1 < find_run2


def test_find_run_replace_no_style_one_run() -> None:
    """Test that replacing a run without style works."""
    document = docx.Document()
    paragraph = document.add_paragraph("")
    paragraph.add_run("Hello, world!")

    find_run = run.FindRun(paragraph, (0, 1), (0, 5))
    find_run.replace("Goodbye")

    assert paragraph.text == "Goodbye, world!"


def test_find_run_replace_twice() -> None:
    """Test that replacing a run twice works."""
    document = docx.Document()
    paragraph = document.add_paragraph("")
    paragraph.add_run("Hello, world!")

    find_run = run.FindRun(paragraph, (0, 1), (0, 5))
    find_run.replace("Goodbye")

    with pytest.raises(ValueError):
        find_run.replace("Goodbye")


def test_find_run_replace_no_style_multiple_runs() -> None:
    """Test that replacing multiple runs without style works."""
    document = docx.Document()
    paragraph = document.add_paragraph("")
    paragraph.add_run("Hello, ")
    paragraph.add_run("world!")

    find_run = run.FindRun(paragraph, (0, 2), (5, 6))
    find_run.replace("Goodbye")

    assert paragraph.text == "HelloGoodbye"


def test_find_run_replace_with_style() -> None:
    """Test that replacing a run with style works."""
    document = docx.Document()
    paragraph = document.add_paragraph("")
    paragraph.add_run("Hello, world!")

    find_run = run.FindRun(paragraph, (0, 1), (0, 5))
    find_run.replace("Goodbye", styles.RunStyle(bold=True))

    assert paragraph.text == "Goodbye, world!"
    assert not paragraph.runs[0].bold
    assert paragraph.runs[1].bold
    assert paragraph.runs[1].text == "Goodbye"


def test_extend_run_format() -> None:
    """Test that formatting a run works."""
    document = docx.Document()
    paragraph = document.add_paragraph("Hello, world!")
    paragraph_run = paragraph.runs[0]

    extend_run = run.ExtendRun(paragraph_run)
    extend_run.format(
        styles.RunStyle(
            bold=True,
            italic=True,
            underline=True,
            superscript=True,
            font_rgb=(1, 0, 0),
        )
    )

    assert paragraph_run.bold
    assert paragraph_run.italic
    assert paragraph_run.underline
    assert paragraph_run.font.superscript
    assert paragraph_run.font.color.rgb == (1, 0, 0)

#!/usr/bin/env python3
"""
DocForge PDF to Word Converter Module
Updated to work with existing DocForge BaseProcessor architecture
"""

import sys
import os
from pathlib import Path
import tempfile
import gc
import time
from typing import Optional, List, Dict, Any
from docx import Document

# Import existing DocForge base class
try:
    from ..core.base import BaseProcessor
except ImportError:
    try:
        from docforge.core.base import BaseProcessor
    except ImportError:
        print("Warning: Could not import BaseProcessor, using fallback")
        from abc import ABC, abstractmethod
        import logging


        class BaseProcessor(ABC):
            def __init__(self, verbose: bool = False):
                self.verbose = verbose
                self.logger = logging.getLogger(self.__class__.__name__)
                if verbose:
                    self.logger.setLevel(logging.DEBUG)

            @abstractmethod
            def process(self, *args, **kwargs) -> Dict[str, Any]:
                pass

            def validate_input(self, input_path: str) -> bool:
                from pathlib import Path
                path = Path(input_path)
                if not path.exists():
                    raise FileNotFoundError(f"Input file not found: {input_path}")
                return True

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_SECTION
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml.parser import parse_xml
    from PyPDF2 import PdfReader
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image

    HAS_DEPENDENCIES = True
except ImportError as e:
    HAS_DEPENDENCIES = False
    missing_lib = str(e).split("'")[1] if "'" in str(e) else str(e)
    print(f"Warning: PDF to Word conversion not available - missing {missing_lib}")


class PDFToWordConverter(BaseProcessor):
    """
    DocForge PDF to Word Converter

    Converts PDF files to Word documents with layout preservation and OCR capabilities.
    Follows existing DocForge BaseProcessor architecture.
    """

    def __init__(self, verbose: bool = False, dpi: int = 200):
        """
        Initialize the PDF to Word converter.

        Args:
            verbose (bool): Enable verbose logging
            dpi (int): DPI for image conversion (default: 200)
        """
        super().__init__(verbose=verbose)
        self.dpi = dpi
        self.batch_size = 3  # Process pages in batches to manage memory
        self.has_dependencies = HAS_DEPENDENCIES

        if not self.has_dependencies:
            if verbose:
                print(
                    "⚠️  PDFToWordConverter: Missing dependencies - install python-docx, PyPDF2, pdf2image, pytesseract, Pillow")

    def process(self, input_path: str, output_path: str, **kwargs) -> Dict[str, Any]:
        """
        Main processing method - converts PDF to Word document.
        Follows DocForge BaseProcessor pattern.

        Args:
            input_path (str): Path to input PDF file
            output_path (str): Path to output Word file
            **kwargs: Additional options (preserve_layout, dpi, etc.)

        Returns:
            Dict[str, Any]: Result dictionary with success status and details
        """
        if not self.has_dependencies:
            return {
                'success': False,
                'error': 'Missing required dependencies for PDF to Word conversion'
            }

        try:
            # Validate input using parent method
            self.validate_input(input_path)

            preserve_layout = kwargs.get('preserve_layout', True)
            dpi = kwargs.get('dpi', self.dpi)

            input_path = Path(input_path)
            output_path = Path(output_path)

            # Create output directory if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)

            if self.verbose:
                print(f"Converting PDF to Word...")
                print(f"Input: {input_path}")
                print(f"Output: {output_path}")
                print(f"Mode: {'Images + Text' if preserve_layout else 'Structured Text Only'}")

            # Extract text and layout from PDF
            page_data = self._extract_text_with_layout(str(input_path), dpi)

            if not page_data:
                return {
                    'success': False,
                    'error': 'No pages extracted from PDF'
                }

            # Create Word document
            self._create_word_document(page_data, str(output_path), preserve_layout)

            if self.verbose:
                print(f"Successfully converted PDF to Word: {output_path}")

            return {
                'success': True,
                'input_file': str(input_path),
                'output_file': str(output_path),
                'pages_processed': len(page_data),
                'preserve_layout': preserve_layout
            }

        except Exception as e:
            error_msg = f"Error during conversion: {e}"
            if self.verbose:
                print(error_msg)
            return {
                'success': False,
                'error': error_msg,
                'input_file': str(input_path) if 'input_path' in locals() else None
            }

    def _extract_text_with_layout(self, pdf_path: str, dpi: int) -> List[Dict[str, Any]]:
        """Extract text from PDF with layout information."""
        if self.verbose:
            print(f"Converting PDF pages to images...")

        all_page_data = []

        # Get total page count
        try:
            import subprocess
            result = subprocess.run(['pdfinfo', str(pdf_path)],
                                    capture_output=True, text=True, timeout=30)
            total_pages = 0
            for line in result.stdout.split('\n'):
                if line.startswith('Pages:'):
                    total_pages = int(line.split(':')[1].strip())
                    break
            if total_pages == 0:
                raise Exception("Could not determine page count")
        except:
            # Fallback method
            images = convert_from_path(pdf_path, dpi=dpi)
            total_pages = len(images)
            all_page_data.extend(self._process_pdf_batch(images, 0, total_pages))
            del images
            gc.collect()
            return all_page_data

        if self.verbose:
            print(f"Found {total_pages} pages - processing in batches of {self.batch_size}")

        # Process in batches
        for batch_start in range(0, total_pages, self.batch_size):
            batch_end = min(batch_start + self.batch_size, total_pages)
            if self.verbose:
                print(f"Processing pages {batch_start + 1}-{batch_end}...")

            images = convert_from_path(
                pdf_path,
                dpi=dpi,
                first_page=batch_start + 1,
                last_page=batch_end
            )

            batch_data = self._process_pdf_batch(images, batch_start, total_pages)
            all_page_data.extend(batch_data)

            del images
            del batch_data
            gc.collect()
            if os.name == 'nt':
                time.sleep(0.5)

        return all_page_data

    def _process_pdf_batch(self, images, batch_start: int, total_pages: int) -> List[Dict[str, Any]]:
        """Process a batch of PDF images with OCR."""
        batch_data = []

        for i, image in enumerate(images):
            page_num = batch_start + i + 1
            if self.verbose:
                print(f"Processing page {page_num}/{total_pages}...")

            try:
                # Use detailed OCR to get word-level positioning
                ocr_data = pytesseract.image_to_data(
                    image,
                    output_type=pytesseract.Output.DICT,
                    config=r'--oem 3 --psm 6'
                )

                # Extract full text
                full_text = pytesseract.image_to_string(image, config=r'--oem 3 --psm 6')

                # Group text by lines based on Y coordinates
                lines = self._extract_text_lines(ocr_data, image.size)

                batch_data.append({
                    'page_num': page_num,
                    'image': image,
                    'full_text': full_text,
                    'lines': lines,
                    'image_size': image.size
                })

                if os.name == 'nt':
                    gc.collect()

            except Exception as e:
                if self.verbose:
                    print(f"Warning: OCR failed for page {page_num}: {e}")
                batch_data.append({
                    'page_num': page_num,
                    'image': image,
                    'full_text': '',
                    'lines': [],
                    'image_size': image.size
                })

        return batch_data

    def _extract_text_lines(self, ocr_data: Dict, image_size: tuple) -> List[Dict[str, Any]]:
        """Extract text organized by lines from OCR data with better positioning."""
        img_width, img_height = image_size
        n_boxes = len(ocr_data['text'])

        # Collect all words with their positions
        words = []
        for i in range(n_boxes):
            text = ocr_data['text'][i].strip()
            conf = int(ocr_data['conf'][i]) if ocr_data['conf'][i] != '-1' else 0
            level = ocr_data['level'][i]

            if text and conf > 40 and level == 5:  # Only confident word-level text
                words.append({
                    'text': text,
                    'left': ocr_data['left'][i],
                    'top': ocr_data['top'][i],
                    'right': ocr_data['left'][i] + ocr_data['width'][i],
                    'bottom': ocr_data['top'][i] + ocr_data['height'][i],
                    'width': ocr_data['width'][i],
                    'height': ocr_data['height'][i],
                    'conf': conf
                })

        if not words:
            return []

        # Sort words primarily by Y position, then by X position
        words.sort(key=lambda w: (w['top'], w['left']))

        # Group words into lines using clustering approach
        lines = []
        current_line = []
        line_height_threshold = 15  # Pixels tolerance for same line

        for word in words:
            if not current_line:
                current_line = [word]
            else:
                avg_top = sum(w['top'] for w in current_line) / len(current_line)

                if abs(word['top'] - avg_top) <= line_height_threshold:
                    current_line.append(word)
                else:
                    if current_line:
                        current_line.sort(key=lambda w: w['left'])
                        line_text = ' '.join(w['text'] for w in current_line)

                        lines.append({
                            'text': line_text,
                            'y_pos': avg_top,
                            'words': current_line.copy(),
                            'left': min(w['left'] for w in current_line),
                            'right': max(w['right'] for w in current_line)
                        })

                    current_line = [word]

        # Add the last line
        if current_line:
            current_line.sort(key=lambda w: w['left'])
            line_text = ' '.join(w['text'] for w in current_line)
            avg_top = sum(w['top'] for w in current_line) / len(current_line)

            lines.append({
                'text': line_text,
                'y_pos': avg_top,
                'words': current_line.copy(),
                'left': min(w['left'] for w in current_line),
                'right': max(w['right'] for w in current_line)
            })

        # Sort lines by Y position
        lines.sort(key=lambda line: line['y_pos'])

        return lines

    def _create_word_document(self, page_data: List[Dict], output_path: str, preserve_images: bool):
        """Create Word document from extracted page data."""
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        for i, page in enumerate(page_data):
            if self.verbose:
                print(f"Creating Word page {page['page_num']}...")

            # Add page break between pages (except first page)
            if i > 0:
                doc.add_page_break()

            if preserve_images and page.get('image'):
                self._add_page_with_image_and_text(doc, page)
            else:
                self._add_formatted_text_page(doc, page)

        # Save document
        doc.save(output_path)
        if self.verbose:
            print(f"Word document saved: {output_path}")

    def _add_page_with_image_and_text(self, doc: Document, page: Dict):
        """Add page with background image and properly formatted text below."""
        temp_fd, temp_path = tempfile.mkstemp(suffix='.png', prefix='docforge_page_')

        try:
            os.close(temp_fd)

            # Convert image if needed
            image = page['image']
            if image.mode in ('RGBA', 'P'):
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background

            # Save image
            image.save(temp_path, 'PNG')

            # Calculate appropriate size
            img_width, img_height = image.size
            page_width = 7.5  # inches
            aspect_ratio = img_height / img_width
            display_width = min(page_width, 7.0)

            # Add paragraph with image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run()

            try:
                run.add_picture(temp_path, width=Inches(display_width))
            except Exception as e:
                if self.verbose:
                    print(f"Could not add image: {e}")
                self._add_formatted_text_page(doc, page)
                return

            # Add editable text content below image
            if page.get('lines'):
                doc.add_paragraph()

                header_para = doc.add_paragraph("EDITABLE TEXT VERSION:")
                header_run = header_para.runs[0]
                header_run.font.bold = True
                header_run.font.size = Pt(12)
                header_run.font.color.rgb = RGBColor(0, 0, 128)

                doc.add_paragraph()

                for line in page['lines']:
                    if line['text'].strip():
                        p = doc.add_paragraph(line['text'].strip())
                        p.style.font.size = Pt(11)
                        p.style.font.name = 'Calibri'

        finally:
            try:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except:
                pass

    def _add_formatted_text_page(self, doc: Document, page: Dict):
        """Add page with professionally formatted text that recreates PDF layout structure."""
        if not page.get('lines'):
            return

        page_num = page['page_num']

        # Add page break between pages (except first page)
        if page_num > 1:
            doc.add_page_break()

        # Process each line with sophisticated formatting
        for i, line in enumerate(page['lines']):
            line_text = line['text'].strip()
            if not line_text:
                continue

            p = doc.add_paragraph()

            # Apply formatting based on content patterns
            self._apply_text_formatting(p, line_text, page['lines'], i)

    def _apply_text_formatting(self, p, line_text: str, all_lines: List, line_index: int):
        """Apply intelligent formatting based on content analysis."""

        # MAIN HEADERS
        if any(header in line_text.upper() for header in [
            'CANADA', 'PROVINCE DE QUEBEC', 'COUR SUPERIEURE', 'COUR DU QUEBEC'
        ]) and len(line_text) < 50:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_text)
            run.font.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run.font.all_caps = True

        # DOCUMENT TITLE
        elif 'PROCES-VERBAL' in line_text.upper():
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_text)
            run.font.bold = True
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'

        # CASE INFORMATION
        elif any(field in line_text for field in ['District :', 'N°', 'Ne :']):
            self._format_field_line(p, line_text)

        # PARTY NAMES AND ROLES
        elif any(role in line_text for role in ['DEMANDE', 'DEFENSE', 'SANTAMARIA', 'ZHANG']):
            if 'ET AL' in line_text:
                p.left_indent = Inches(0.5)
            run = p.add_run(line_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

        # DATE AND PRESENT
        elif line_text.startswith(('DATE:', 'PRESENT:')):
            self._format_field_line(p, line_text)

        # LAWYER INFORMATION
        elif line_text.startswith('Me ') or any(law in line_text for law in ['avocats', 'RPGL', 'S.E.N.C.R.L']):
            p.left_indent = Inches(1.0)
            run = p.add_run(line_text)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

        # SECTION HEADERS
        elif line_text.upper() in ['REFERENCES', 'PLAIDOIRIES', 'JUGEMENT']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_text)
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            p.space_before = Pt(12)

        # TIME STAMPS
        elif self._is_timestamp(line_text):
            self._format_timestamp(p, line_text)

        # Q&A SECTIONS
        elif line_text.strip().startswith(('Q.', 'A.')) or line_text.strip().endswith('?'):
            self._format_qa_line(p, line_text)

        # CONSIDERANT clauses
        elif line_text.startswith('CONSIDERANT'):
            p.left_indent = Inches(0.0)
            p.space_before = Pt(6)
            run = p.add_run(line_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

        # SIGNATURES AND OFFICIALS
        elif any(sig in line_text.upper() for sig in [
            'HONORABLE', 'J.C.S.', 'G.A.C.S.', 'MAXIME CHABOT', 'RENEE THERIAULT'
        ]):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(12)
            run = p.add_run(line_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

        # DEFAULT FORMATTING
        else:
            if (line_index > 0 and not line_text[0].isupper() and len(line_text) > 20):
                p.left_indent = Inches(0.5)

            run = p.add_run(line_text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    def _format_field_line(self, p, line_text: str):
        """Format field lines with bold labels."""
        run = p.add_run(line_text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

        if ':' in line_text:
            parts = line_text.split(':', 1)
            p.clear()
            label_run = p.add_run(parts[0] + ':')
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            if len(parts) > 1:
                value_run = p.add_run(' ' + parts[1].strip())
                value_run.font.size = Pt(11)

    def _is_timestamp(self, line_text: str) -> bool:
        """Check if line contains a timestamp."""
        return (line_text[:5].replace('h', '').replace(':', '').isdigit() or
                any(time in line_text[:6] for time in ['09h', '10h', '11h', '12h', '13h', '14h', '15h', '16h']))

    def _format_timestamp(self, p, line_text: str):
        """Format timestamp lines."""
        p.left_indent = Inches(0.3)
        parts = line_text.split(' ', 1)

        if len(parts) >= 2 and 'h' in parts[0]:
            time_run = p.add_run(parts[0])
            time_run.font.bold = True
            time_run.font.size = Pt(10)
            time_run.font.name = 'Times New Roman'

            desc_run = p.add_run(' ' + ' '.join(parts[1:]))
            desc_run.font.size = Pt(10)
            desc_run.font.name = 'Times New Roman'
        else:
            run = p.add_run(line_text)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    def _format_qa_line(self, p, line_text: str):
        """Format Q&A lines."""
        p.left_indent = Inches(0.8)
        p.first_line_indent = Inches(-0.3)

        if line_text.startswith(('Q.', 'A.')):
            qa_marker = line_text[:2]
            rest_text = line_text[2:].strip()

            marker_run = p.add_run(qa_marker)
            marker_run.font.bold = True
            marker_run.font.size = Pt(10)
            marker_run.font.name = 'Times New Roman'

            text_run = p.add_run(' ' + rest_text)
            text_run.font.size = Pt(10)
            text_run.font.name = 'Times New Roman'
            text_run.font.italic = True
        else:
            run = p.add_run(line_text)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.font.italic = True


# Backward compatibility functions for your original API
def check_dependencies():
    """Check if required dependencies are available."""
    return HAS_DEPENDENCIES


def pdf_to_word(pdf_path: str, word_path: str, preserve_layout: bool = True, dpi: int = 200):
    """
    Convert PDF to Word document - backward compatibility function.

    Args:
        pdf_path (str): Input PDF file path
        word_path (str): Output Word file path
        preserve_layout (bool): Whether to preserve visual layout
        dpi (int): DPI for image conversion

    Returns:
        bool: True if successful
    """
    converter = PDFToWordConverter(verbose=True, dpi=dpi)
    result = converter.process(pdf_path, word_path,
                               preserve_layout=preserve_layout, dpi=dpi)
    return result.get('success', False)


# Command line interface for standalone usage
def main():
    """Command line interface - preserved from original."""
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word.py <input_pdf> <output_word> [preserve_layout] [dpi]")
        print("Example: python pdf_to_word.py document.pdf document.docx true 200")
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_word = sys.argv[2]
    preserve_layout = sys.argv[3].lower() == 'true' if len(sys.argv) > 3 else True
    dpi = int(sys.argv[4]) if len(sys.argv) > 4 else 200

    success = pdf_to_word(input_pdf, output_word, preserve_layout=preserve_layout, dpi=dpi)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
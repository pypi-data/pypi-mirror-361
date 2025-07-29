from docx import Document
from docxcompose.composer import Composer

def chk_br(par):
    for run in par.r_lst:
        if len(run.br_lst) > 0:
            return(True)
    return(False)

def doc_extract(doc_directory,fab_filename,skipped_pg_br,keeped_pg_br):
    tmp_file = doc_directory + "/schemas.docx"

    doc_fab = Document(fab_filename)
    doc_fab.save(tmp_file)
    doc_tmp = Document(tmp_file)

    # extraction des pages du docuemnt fabriquant
    num_page_breaks = 0

    state = 0

    for elt_idx, elt in enumerate(doc_tmp.element.body):
        if state == 0:
            if elt.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
                if chk_br(elt):
                    num_page_breaks += 1
                if num_page_breaks == 2:
                    doc_tmp.save(tmp_file)
                    state = 1
                doc_tmp.element.body.remove(elt)
        elif state == 1:
            if elt.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
                if chk_br(elt):
                    num_page_breaks += 1
                if num_page_breaks == 3:
                    state = 2
                    doc_tmp.element.body.remove(elt)
                #doc_tmp.add_paragraph()
            #elif elt.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl":
                #table_copy(elt, doc_fab, doc_tmp)
        elif state == 2:
            if elt.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr":
                doc_tmp.element.body.remove(elt)
            
    doc_tmp.save(tmp_file)

    return(tmp_file)

def compose_doc(base_file,files_list,final_doc):
    files_list = files_list.split(";")
    master = Document(base_file)
    composer = Composer(master)
    for filename in files_list:
        doc_temp = Document(filename)
        composer.append(doc_temp)
    composer.save(final_doc)
    return

if __name__ == '__main__':
    flist = ['//winprod/commun/is/chantier/1000/schemas.docx', '//winprod/commun/is/chantier/1000/1000_DETAIL.docx', '//winprod/commun/is/chantier/_doc_base/Formulaire_ALU_cgv.docx']
    compose_doc('//winprod/commun/is/chantier/_doc_base/Formulaire_ALU_1-3.docx',flist,'//winprod/commun/is/chantier/1000/Devis.docx')

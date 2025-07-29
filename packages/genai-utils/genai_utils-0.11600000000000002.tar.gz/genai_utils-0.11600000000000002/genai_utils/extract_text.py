#!/usr/bin/env python

import sys, os, logging, pandas as pd
logger = logging.getLogger( "genai_utils" )
from mangorest.mango import webapi

# ------------------------------------------------------------------------------------------
# Simply extarct text rfrom PDF file
def extractTextPDF(file):
    import pdfplumber

    #from pdfplumber.page import Page
    #from pdfplumber.pdf import PDF
    #from pdfplumber.table import Table
    #from pdfplumber.utils import intersects_bbox

    text = []
    #print(f"**** OPENINF FILE ===>  {f}")
    with pdfplumber.open(file) as doc:
        for page in doc.pages:
            #lines = page.extract_text_lines()
            #txt = "\n".join([l['text'] for l in lines])
            txt = page.extract_text_simple()
            text.append(txt)

    all="\n".join(text)
    return all

# ------------------------------------------------------------------------------------------
# Simply extarct text rfrom PDF file
def extractDocx(file):
    import docx
    document = docx.Document(file)
    txts=[]
    for p in document.paragraphs:
        txts.append(p.text)

    all = "\n".join(txts)
    return all
#-----------------------------------------------------------------------------------------    
@webapi("/gpt/extractText/")
def extractText(request=None, file=None, **kwargs):
    ret = f"Unknown file type {file}"

    if ( request and not file):
        for f in request.FILES.getlist('file'):
            content = f.read()
            #fileIO = io.BytesIO(content)
            file = f"/tmp/{str(f)}"
            with open(file, "wb") as f:
                f.write(content)

    print(f"Reading file {file}")

    if (file.endswith(".doc") or file.endswith(".docx") ):
        ret =  extractDocx(file)
    elif (file.endswith(".txt") or file.endswith(".md") ):
        with open(file, "r", encoding="utf-8", errors='ignore') as f:
            ret = f.read()
    elif (file.endswith(".pdf") ):
        ret = extractTextPDF(file)
    elif (file.endswith(".xlsx") or file.endswith(".xls")):
        df = pd.read_excel(file)
        ret = df.to_html()
    elif file.endswith(".csv") :
        df = pd.read_csv(file)
        ret = df.to_html()
    elif file:
        ret = open(file, "rb").read()
    else:
        ret = ""
    
    return ret

# ---------------------------------------------------------------------------------------
def getChunks(file, chunk_size=2000, overlap=256 ):
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    
    txt= extractText( file= file)

    split = RecursiveCharacterTextSplitter(
        chunk_size= chunk_size,
        chunk_overlap=overlap,
        length_function=len,
        add_start_index=True,
    )
    docs = []
    for txt in split.split_text(txt):
        d = Document(page_content= txt,  metadata=dict(source= file))
        docs.append(d)

    return docs

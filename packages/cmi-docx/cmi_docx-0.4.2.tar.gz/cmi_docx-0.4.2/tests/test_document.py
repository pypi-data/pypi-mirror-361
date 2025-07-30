"""Tests for the document module."""

import docx
import pytest

from cmi_docx import document, styles


def test_find_in_paragraphs() -> None:
    """Test finding a text in a document's paragraphs."""
    doc = docx.Document()
    doc.add_paragraph("Hello, world!")
    doc.add_paragraph("Hello, world, Hello!")
    extend_document = document.ExtendDocument(doc)

    actual = extend_document.find_in_paragraphs("Hello")

    assert actual[0].character_indices == [(0, 5)]
    assert actual[1].character_indices == [(0, 5), (14, 19)]


def test_find_in_runs() -> None:
    """Test finding a text in a document's runs."""
    doc = docx.Document()
    paragraph = doc.add_paragraph("Hello, world!")
    paragraph.add_run("Hello, world, Hello!")
    extend_document = document.ExtendDocument(doc)

    actual = extend_document.find_in_runs("Hello")

    assert actual[0].run_indices == (0, 0)
    assert actual[0].character_indices == (0, 5)
    assert actual[1].run_indices == (1, 1)
    assert actual[1].character_indices == (0, 5)
    assert actual[2].run_indices == (1, 1)
    assert actual[2].character_indices == (14, 19)


@pytest.mark.parametrize(
    ("runs", "needle", "replace", "expected"),
    [
        (
            ["Hello, world!"],
            "Hello",
            "Goodbye",
            "Goodbye, world!",
        ),
        (
            ["Hello, world!", " Hello!"],
            "Hello",
            "Goodbye",
            "Goodbye, world! Goodbye!",
        ),
        (
            ["Hello {{", "FULL_NAME}}"],
            "{{FULL_NAME}}",
            'Shizuka "Lea" Sakai',
            'Hello Shizuka "Lea" Sakai',
        ),
        (
            ["This is James", " Bond ", "007!"],
            "James Bond 007",
            "Patrick",
            "This is Patrick!",
        ),
        (
            ["This is Alec", " Travelyan ", "006!"],
            "This is Alec Travelyan 006!",
            "",
            "",
        ),
        (
            ["This", " is ", "Patrick!"],
            "",
            "Nonsense",
            "This is Patrick!",
        ),
    ],
)
def test_replace(runs: list[str], needle: str, replace: str, expected: str) -> None:
    """Test replacing text in a document."""
    doc = docx.Document()
    paragraph = doc.add_paragraph(runs[0])
    for run in runs[1:]:
        paragraph.add_run(run)
    extend_document = document.ExtendDocument(doc)

    extend_document.replace(needle, replace)

    assert doc.paragraphs[0].text == expected


def test_replace_with_style() -> None:
    """Test replacing text in a document with style."""
    doc = docx.Document()
    paragraph = doc.add_paragraph("")
    paragraph.add_run("{{")
    paragraph.add_run("Hello, World!")
    paragraph.add_run("}}")
    extend_document = document.ExtendDocument(doc)

    extend_document.replace("{{Hello", "Goodbye", styles.RunStyle(bold=True))

    assert doc.paragraphs[0].text == "Goodbye, World!}}"
    assert not doc.paragraphs[0].runs[0].bold
    assert doc.paragraphs[0].runs[1].bold
    assert doc.paragraphs[0].runs[1].text == "Goodbye"


def test_insert_paragraph_by_object() -> None:
    """Test inserting a paragraph into a document."""
    doc = docx.Document()
    doc.add_paragraph("Hello, world!")
    doc.add_paragraph("Goodbye, world!")
    extend_document = document.ExtendDocument(doc)
    new_paragraph = docx.Document().add_paragraph("Maintain, world!")

    extend_document.insert_paragraph_by_object(1, new_paragraph)

    assert doc.paragraphs[0].text == "Hello, world!"
    assert doc.paragraphs[1].text == "Maintain, world!"
    assert doc.paragraphs[2].text == "Goodbye, world!"


def test_insert_paragraph_by_text() -> None:
    """Test inserting a paragraph into a document."""
    doc = docx.Document()
    doc.add_paragraph("Hello, world!")
    doc.add_paragraph("Goodbye, world!")
    extend_document = document.ExtendDocument(doc)

    extend_document.insert_paragraph_by_text(1, "Maintain, world!")

    assert doc.paragraphs[0].text == "Hello, world!"
    assert doc.paragraphs[1].text == "Maintain, world!"
    assert doc.paragraphs[2].text == "Goodbye, world!"

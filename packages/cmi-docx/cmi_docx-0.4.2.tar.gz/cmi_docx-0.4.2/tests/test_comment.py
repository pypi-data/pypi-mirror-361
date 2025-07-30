"""Tests for the comment module."""

import docx
from docx.opc import constants as docx_constants

import cmi_docx
from cmi_docx import comment


def test_add_comment_single() -> None:
    """Tests adding a comment to a single entity."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    author = "Grievous"
    message = "Ah, General Kenobi."

    cmi_docx.add_comment(document, para, author, message)
    para_comment = para.part.part_related_by(docx_constants.RELATIONSHIP_TYPE.COMMENTS)

    assert author in para_comment.blob.decode()
    assert message in para_comment.blob.decode()


def test_add_comment_range() -> None:
    """Tests adding a comment with a range of entities."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    run_start = para.add_run(" run 3!")
    run_end = para.add_run(" run 4!")
    author = "Grievous"
    message = "Ah, General Kenobi."

    cmi_docx.add_comment(document, (run_start, run_end), author, message)
    comment_start = run_start.part.part_related_by(
        docx_constants.RELATIONSHIP_TYPE.COMMENTS
    )
    comment_end = run_end.part.part_related_by(
        docx_constants.RELATIONSHIP_TYPE.COMMENTS
    )

    assert author in comment_start.blob.decode()
    assert message in comment_start.blob.decode()
    assert author in comment_end.blob.decode()
    assert message in comment_end.blob.decode()


def test_comment_preserver_extract_comments() -> None:
    """Tests getting comments from a paragraph."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    author = "Grievous"
    message = "Ah, General Kenobi."
    cmi_docx.add_comment(document, para, author, message)
    preserver = comment.CommentPreserver(para._element)

    comments = preserver.extract_comments()

    assert len(comments) == 1
    assert comments[0].start_index == 0
    assert comments[0].end_index == len(para.text)


def test_comment_preserver_strip_comments() -> None:
    """Tests stripping comments from a paragraph."""
    document = docx.Document()
    para = document.add_paragraph()
    run = para.add_run("this")
    para.add_run("is")
    cmi_docx.add_comment(document, run, "Grievous", "Ah, General Kenobi.")
    preserver = comment.CommentPreserver(para._element)

    preserver.strip_comments()

    assert para.text == "thisis"


def test_replace_between_preserve_comments_contained() -> None:
    """Using the replace_between with an edit inside the comment."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    extend_para = cmi_docx.ExtendParagraph(para)
    author = "Grievous"
    message = "Ah, General Kenobi."
    cmi_docx.add_comment(document, para, author, message)

    extend_para.replace_between(5, 7, "was")
    comments = comment.CommentPreserver(para._element).extract_comments()

    assert para.text == "This was a sample paragraph."
    assert len(comments) == 1
    assert comments[0].start_index == 0
    assert comments[0].end_index == len(para.text)


def test_strip_comments_removes_comments_entire_string() -> None:
    """Using the replace_between with the edit equivalent to the comment range."""
    document = docx.Document()
    para = document.add_paragraph("")
    run = para.add_run("this")
    para.add_run("is")
    extend_para = cmi_docx.ExtendParagraph(para)
    cmi_docx.add_comment(document, run, "Grievous", "Ah, General Kenobi.")

    extend_para.replace_between(0, len(run.text), "was")
    comments = comment.CommentPreserver(para._element).extract_comments()

    assert para.text == "wasis"
    assert len(comments) == 1
    assert comments[0].start_index == 0
    assert comments[0].end_index == len(para.runs[0].text)


def test_restore_comments_preserves_comment_references() -> None:
    """Tests that restoring comments doesn't remove comment reference runs."""
    document = docx.Document()
    para = document.add_paragraph("Sample text with comment.")
    author = "Obi-Wan"
    message = "Hello there!"

    cmi_docx.add_comment(document, para, author, message)
    cmi_docx.add_comment(document, para, author, message)

    original_xml = para._element
    comment_refs_before = len(original_xml.xpath(".//w:commentReference"))

    preserver = comment.CommentPreserver(para._element)
    extracted_comments = preserver.extract_comments()
    preserver.restore_comments(extracted_comments)
    comment_refs_after = len(para._element.xpath(".//w:commentReference"))

    assert comment_refs_before > 0
    assert comment_refs_after == comment_refs_before
    assert para.text == "Sample text with comment."

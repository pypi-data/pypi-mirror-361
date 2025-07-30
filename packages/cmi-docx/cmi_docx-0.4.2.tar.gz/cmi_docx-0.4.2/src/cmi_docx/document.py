"""Extends a python-docx Word document with additional functionality."""

import pathlib

from docx import document
from docx.text import paragraph as docx_paragraph

from cmi_docx import paragraph, run, styles


class ExtendDocument:
    """Extends a python-docx Word document with additional functionality."""

    def __init__(self, document: document.Document) -> None:
        """Initializes a DocxSearch object for finding text."""
        self.document = document

    def find_in_paragraphs(self, needle: str) -> list[paragraph.FindParagraph]:
        """Finds the indices of a text relative to the paragraphs.

        Args:
            needle: The text to find.

        Returns:
            The indices of the text in the document.
        """
        return [
            paragraph.ExtendParagraph(para).find_in_paragraph(needle)
            for para in self.all_paragraphs
        ]

    def find_in_runs(self, needle: str) -> list[run.FindRun]:
        """Finds the indices of a text relative to the document's runs.

        Args:
            needle: The text to find.

        Returns:
            The locations of the text in the runs.
        """
        return [
            finder
            for para in self.all_paragraphs
            for finder in paragraph.ExtendParagraph(para).find_in_runs(needle)
        ]

    def replace(
        self, needle: str, replace: str, style: styles.RunStyle | None = None
    ) -> None:
        """Finds and replaces text in a Word document.

        Args:
            needle: The text to find.
            replace: The text to replace.
            style: The style to apply to the replacement text.

        """
        run_finder = self.find_in_runs(needle)
        run_finder.sort(
            key=lambda x: (x.run_indices[0], x.character_indices[0]), reverse=True
        )

        for run_find in run_finder:
            run_find.replace(replace, style)

    def insert_paragraph_by_text(
        self,
        index: int,
        text: str,
        style: str | None = None,
    ) -> docx_paragraph.Paragraph:
        """Inserts a paragraph into a document.

        Args:
            index: The index to insert the paragraph at.
            text: The text to insert.
            style: The style to apply to the text.

        Returns:
            The new paragraph.
        """
        new_paragraph = self._insert_empty_paragraph(index, style)
        new_paragraph.add_run(text)
        return new_paragraph

    def insert_paragraph_by_object(
        self,
        index: int,
        paragraph: docx_paragraph.Paragraph,
    ) -> docx_paragraph.Paragraph:
        """Inserts a paragraph into a document.

        Args:
            index: The index to insert the paragraph at.
            paragraph: The paragraph to insert.

        Returns:
            The new paragraph.
        """
        new_paragraph = self._insert_empty_paragraph(index)
        for paragraph_run in paragraph.runs:
            new_paragraph.add_run(paragraph_run.text)
        return new_paragraph

    def insert_image(
        self,
        index: int,
        image_path: str | pathlib.Path,
        width: int | None = None,
        height: int | None = None,
    ) -> docx_paragraph.Paragraph:
        """Inserts an image at a given paragraph index.

        Args:
            index: The paragraph index to insert the image at.
            image_path: The path to the image to insert.
            width: The width of the image.
            height: The height of the image.
        """
        new_paragraph = self._insert_empty_paragraph(index)
        run = new_paragraph.add_run()
        run.add_picture(str(image_path), width=width, height=height)
        return new_paragraph

    @property
    def all_paragraphs(self) -> list[docx_paragraph.Paragraph]:
        """Returns all paragraphs in the document, including headers and footers."""
        all_paragraphs = list(self.document.paragraphs)

        for section in self.document.sections:
            all_paragraphs.extend(
                (*section.footer.paragraphs, *section.header.paragraphs)
            )
        return all_paragraphs

    def _insert_empty_paragraph(
        self, index: int, style: str | None = None
    ) -> docx_paragraph.Paragraph:
        """Inserts an empty paragraph at a given index.

        Args:
            index: The index to insert the paragraph at.
            style: The style to apply to the paragraph.

        Returns:
            The new paragraph.
        """
        n_paragraphs = len(self.document.paragraphs)
        if index > n_paragraphs:
            raise ValueError(f"Index {index} is out of range.")

        if index == n_paragraphs:
            new_paragraph = self.document.add_paragraph(style=style)
        else:
            new_paragraph = new_paragraph = self.document.paragraphs[
                index
            ]._insert_paragraph_before()
            new_paragraph.style = style  # type: ignore[assignment] # Mypy ignores setter types.

        return new_paragraph

from docx import Document
from docx.shared import Pt
import csv

def hello_world():
    print('Hello world fron in_projet')

def build_projet(txt_filename,tpl_filename,doc_filename):
    doc_prj = Document(tpl_filename)
    font_size = Pt(8)

    ret = False
    table = doc_prj.tables[0]

    data = []
    # Ouvrir le fichier en mode lecture
    with open(txt_filename, 'r', encoding="utf-8") as file:
        # Créer un lecteur CSV
        reader = csv.reader(file, delimiter=';')  # Spécifiez le délimiteur approprié si nécessaire
        for rowidx, row in enumerate(reader):
             if rowidx >= 11:
                if len(row) > 0:
                    data.append(row)

        current_pos = ''
        for posidx, pos in enumerate(data[0]):
            if pos != "X":
                current_pos = data[1][posidx]
                cells = table.add_row().cells
                cells[0].text = data[1][posidx]
                cells[0].paragraphs[0].runs[0].font.bold = True
                text = str(data[4][posidx])
                cells[1].text = text.replace("|","\r")
                cells[2].text = data[3][posidx]
                cells[3].text = data[13][posidx]
                cells[4].text = data[5][posidx]
                cells[5].text = data[11][posidx]
                cells[6].text = data[12][posidx]

            #else:
            #    print("desigantion = ",data[4][posidx])

        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = font_size
                    run.font.name = 'Calibri'
        
        doc_prj.save(doc_filename)
        ret = True
    return(ret)

if __name__ == '__main__':
    arguments = sys.argv


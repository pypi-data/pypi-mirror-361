from docx import Document

def chk_br(par):
    for run in par.r_lst:
        if len(run.br_lst) > 0:
            return(True)
    return(False)

def build_word_file(doc_directory,fab_filename,skipped_pg_br,keeped_pg_br):
    tmp_file = doc_directory + "/schemas.docx"

    print(fab_filename)
    doc_fab = Document(fab_filename)
    doc_fab.save(tmp_file)
    doc_tmp = Document(tmp_file)

    # extraction des pages du docuemnt fabriquant
    num_page_breaks = 0

    state = 0

    for elt_idx, elt in enumerate(doc_tmp.element.body):
        print(elt.tag)
        if state == 0:
            if elt.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
                if chk_br(elt):
                    num_page_breaks += 1
                if num_page_breaks == 2:
                    doc_tmp.save(tmp_file)
                    state = 1
                doc_tmp.element.body.remove(elt)
        elif state == 1:
            if elt.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
                if chk_br(elt):
                    num_page_breaks += 1
                if num_page_breaks == 3:
                    state = 2
                    doc_tmp.element.body.remove(elt)
                #doc_tmp.add_paragraph()
            #elif elt.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl":
                #table_copy(elt, doc_fab, doc_tmp)
        elif state == 2:
            if elt.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr":
                doc_tmp.element.body.remove(elt)
            
    print("save " + tmp_file)
    doc_tmp.save(tmp_file)

    return(tmp_file)

def build_devis(doc_directory,fab_filename,startdoc,end_doc):
    tmp_doc = build_word_file(doc_directory,fab_filename,2,1)
    print(tmp_doc)

doc_dir = "Q:/is/chantier/1000/"
fab_fn = "Q:/is/chantier/1000/FABRIQUANT/1000.docx"

build_devis(doc_dir,fab_fn,"","")

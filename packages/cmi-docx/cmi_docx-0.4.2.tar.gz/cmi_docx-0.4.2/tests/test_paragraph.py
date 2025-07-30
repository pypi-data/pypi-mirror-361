"""Tests for the paragraph module."""

import docx
import pytest
from docx.text import paragraph as docx_paragraph

from cmi_docx import paragraph, styles


@pytest.fixture
def sample_paragraph() -> docx_paragraph.Paragraph:
    """Returns a sample paragraph."""
    document = docx.Document()
    return document.add_paragraph("This is a sample paragraph.")


def test_find_single_in_paragraph(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test finding a text in a paragraph."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)
    expected = paragraph.FindParagraph(
        paragraph=sample_paragraph,
        character_indices=[(10, 16)],
    )

    actual = extend_paragraph.find_in_paragraph("sample")

    assert actual == expected


def test_find_multiple_in_paragraph(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test finding multiple texts in a paragraph."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)
    expected = paragraph.FindParagraph(
        paragraph=sample_paragraph,
        character_indices=[(2, 4), (5, 7)],
    )

    actual = extend_paragraph.find_in_paragraph("is")

    assert actual == expected


def test_find_in_single_run(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test finding a text in a single paragraph run."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)
    expected = [
        paragraph.run.FindRun(
            paragraph=sample_paragraph,
            run_indices=(0, 0),
            character_indices=(2, 4),
        ),
        paragraph.run.FindRun(
            paragraph=sample_paragraph,
            run_indices=(0, 0),
            character_indices=(5, 7),
        ),
    ]

    actual = extend_paragraph.find_in_runs("is")

    assert actual[0].paragraph.text == expected[0].paragraph.text
    assert actual[0].run_indices == expected[0].run_indices
    assert actual[0].character_indices == expected[0].character_indices
    assert actual[1].paragraph.text == expected[1].paragraph.text
    assert actual[1].run_indices == expected[1].run_indices
    assert actual[1].character_indices == expected[1].character_indices


def test_find_in_single_run_complete() -> None:
    """Tests finding an exact match of a run."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    para.add_run("full-run")
    para.add_run("another")
    extend_paragraph = paragraph.ExtendParagraph(para)

    actual = extend_paragraph.find_in_runs("full-run")

    assert len(actual[0].runs) == 1


def test_replace_single_run(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test replacing text in a paragraph."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)

    extend_paragraph.replace("sample", "example")

    assert sample_paragraph.text == "This is a example paragraph."


def test_replace_multiple_runs(
    sample_paragraph: docx_paragraph.Paragraph,
) -> None:
    """Test replacing text in multiple runs."""
    sample_paragraph.add_run(" This is a sample paragraph.")
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)

    extend_paragraph.replace("This is", "That was")

    assert (
        sample_paragraph.text
        == "That was a sample paragraph. That was a sample paragraph."
    )


def test_insert_run_middle() -> None:
    """Test inserting a run into a paragraph."""
    document = docx.Document()
    para = document.add_paragraph("")
    para.add_run("Hello ")
    para.add_run("world!")
    extend_paragraph = paragraph.ExtendParagraph(para)

    extend_paragraph.insert_run(1, "beautiful ", styles.RunStyle(bold=True))

    assert para.text == "Hello beautiful world!"
    assert para.runs[1].bold


def test_insert_run_start() -> None:
    """Test inserting a run at the start of a paragraph."""
    document = docx.Document()
    para = document.add_paragraph("world!")
    extend_paragraph = paragraph.ExtendParagraph(para)

    extend_paragraph.insert_run(0, "Hello ", styles.RunStyle(bold=True))

    assert para.text == "Hello world!"
    assert para.runs[0].bold


@pytest.mark.parametrize("index", [-1, 1])
def test_insert_run_end(index: int) -> None:
    """Test inserting a run at the end of a paragraph."""
    document = docx.Document()
    para = document.add_paragraph("Hello")
    extend_paragraph = paragraph.ExtendParagraph(para)

    extend_paragraph.insert_run(index, " world!", styles.RunStyle(bold=True))

    assert para.text == "Hello world!"
    assert para.runs[1].bold


def test_insert_run_empty() -> None:
    """Test inserting a run into an empty paragraph."""
    document = docx.Document()
    para = document.add_paragraph("")
    extend_paragraph = paragraph.ExtendParagraph(para)

    extend_paragraph.insert_run(0, "Hello", styles.RunStyle(bold=True))

    assert para.text == "Hello"
    assert para.runs[0].bold


def test_replace_between_one_run() -> None:
    """Test replacing text in one run."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    extend_paragraph = paragraph.ExtendParagraph(para)

    extend_paragraph.replace_between(5, 7, "was")

    assert para.text == "This was a sample paragraph."
    assert para.runs[0].text == "This "
    assert para.runs[1].text == "was"
    assert para.runs[2].text == " a sample paragraph."


def test_replace_between_multiple_runs() -> None:
    """Test replacing text in multiple runs."""
    document = docx.Document()
    para = document.add_paragraph("This ")
    para.add_run("is")
    para.add_run(" Sparta!")
    extend_paragraph = paragraph.ExtendParagraph(para)

    extend_paragraph.replace_between(3, 10, "nk sm")

    assert para.text == "Think smarta!"
    assert para.runs[0].text == "Thi"
    assert para.runs[1].text == "nk sm"
    assert para.runs[2].text == "arta!"

# Empire Chain File Reader Module
# Updated: March 2025 - Adding comments for version tracking

from typing import Protocol
from pathlib import Path
import PyPDF2
import docx
import json
import csv
import requests
import webbrowser
import io
import os

class FileReader(Protocol):
    def read(self, file_path: str) -> str:
        pass

class PDFReader(FileReader):
    def read(self, file_path: str) -> str:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

class DocxReader(FileReader):
    def read(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

class TxtReader(FileReader):
    def read(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

class JSONReader(FileReader):
    def read(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            return json.dumps(data, indent=2)

class CSVReader(FileReader):
    def read(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'r', encoding='utf-8') as file:
            reader = csv.reader(file)
            for row in reader:
                text += ",".join(row) + "\n"
        return text

class GoogleDocsReader(FileReader):
    def read(self, file_path: str) -> str:
        """Reads a Google Drive file and returns its content as text.
        
        Args:
            file_path: The Google Drive file URL
            
        Returns:
            str: The document content as text
            
        Raises:
            ValueError: If file cannot be accessed
        """
        if not 'drive.google.com' in file_path:
            raise ValueError("Not a valid Google Drive URL")
            
        if '/file/d/' in file_path:
            file_id = file_path.split('/file/d/')[1].split('/')[0]
        elif '/document/d/' in file_path:
            file_id = file_path.split('/document/d/')[1].split('/')[0]
        else:
            raise ValueError("Invalid Google Drive URL format. Please use the 'Share' link from Google Drive.")

        download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        
        try:
            response = requests.get(download_url)
            
            if response.status_code == 403 or 'Sign in' in response.text:
                print("\nPlease sign in with your Google account to access this file.")
                print("A browser window will open. After signing in, please try again.")
                webbrowser.open(file_path)
                raise ValueError("Please authenticate through your browser and try again")
                
            if response.status_code != 200:
                raise ValueError("Could not access file. Make sure the file is shared and accessible.")
            
            content = io.BytesIO(response.content)
            
            if b'%PDF' in response.content[:1024]:
                reader = PyPDF2.PdfReader(content)
                return "\n".join(page.extract_text() for page in reader.pages)
            
            try:
                return response.content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    doc = docx.Document(content)
                    return "\n".join(paragraph.text for paragraph in doc.paragraphs)
                except:
                    raise ValueError("Unsupported file type or file is corrupted")
                    
        except requests.RequestException as e:
            raise ValueError(f"Error accessing Google Drive file: {str(e)}")

class DocumentReader:
    def __init__(self):
        """Initialize document reader."""
        self.readers = {
            '.pdf': PDFReader(),
            '.docx': DocxReader(),
            '.txt': TxtReader(),
            '.json': JSONReader(),
            '.csv': CSVReader()
        }
        self.google_reader = GoogleDocsReader()
    
    def _is_google_drive_url(self, file_path: str) -> bool:
        """Check if the given path is a Google Drive URL."""
        return 'drive.google.com' in file_path
    
    def read(self, file_path: str) -> str:
        """Read content from various file types and return as text.
        
        Args:
            file_path: Path to the file to read or Google Drive URL
            
        Returns:
            str: Text content of the file
            
        Raises:
            ValueError: If file type is not supported or file cannot be accessed
        """
        if self._is_google_drive_url(file_path):
            return self.google_reader.read(file_path)
            
        file_extension = Path(file_path).suffix.lower()
        if not file_extension:
            raise ValueError("File has no extension and is not a Google Drive URL")
            
        if file_extension not in self.readers:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return self.readers[file_extension].read(file_path)
    
    def supported_formats(self) -> list[str]:
        """Get list of supported file formats.
        
        Returns:
            list[str]: List of supported file extensions
        """
        return list(self.readers.keys()) 
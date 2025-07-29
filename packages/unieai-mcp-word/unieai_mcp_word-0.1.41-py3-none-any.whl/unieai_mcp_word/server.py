from fastmcp import FastMCP
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import tempfile
from datetime import datetime
import os, shutil, re, uuid

#PATTERN = re.compile(r"{{\s*(\w+)\s*}}")  # 捕捉 {{ key }} 或 {{key}}
#PATTERN = re.compile(r'\(\w+)\') 
#PATTERN = re.compile(r"\b(\w+)\b")

PATTERN = re.compile(r'\$\{(\w+)\}')

def main():
    mcp = FastMCP("unieai-mcp-word-stdio")

    @mcp.tool()
    def write_data_to_word_with_custom(data: dict) -> str:
        """
        UnieAI 專用客戶回饋報告 Word 模板
        將 data 的 key/value 套入 Word 範本中的 ${key} 位置，
        完成後回傳下載超連結（Markdown 內嵌）。
        """
        outputpath = (
            #"/app/data/storage/unieai-mcp-word/customer_feedback_report_"
            "D:/customer_feedback_report_"
            + datetime.now().strftime("%Y%m%d")
            + "_"
            + str(uuid.uuid4())
            + ".docx"
        )
        params = {
            #"filepath": "/app/data/storage/unieai-mcp-word/customer_feedback_report_temp.docx",
            "filepath": "D:/customer_feedback_report_temp.docx",
            "outputpath": outputpath,
            "data_map": data,
            "desc": "客戶回饋報告 Word檔案下載"
        }
        return fill_word_with_context(params)



    mcp.run(transport="stdio")



def replace_in_runs(runs, ctx, test_log):

    for run in runs:
        if PATTERN.search(run.text):
            test_log.append(f"run.text: {run.text}")
            new_text = PATTERN.sub(lambda m: str(ctx.get(m.group(1), m.group(0))), run.text)
            test_log.append(f"new_text: {new_text}")
            run.text = new_text

            # 創建新的 run，將新的文字和格式應用於該 run
            new_run = run._element.getparent().add_run(new_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size

            # 刪除原有的 run
            run._element.getparent().remove(run._element)

def fill_word_with_context(params: dict) -> str:
    src, dst, ctx = params["filepath"], params["outputpath"], params["data_map"]
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    shutil.copy(src, dst)
    doc = Document(dst)

    test_log = []
    test_log.append(f"init_new1: ")

    # 替換段落中的佔位符
    for para in doc.paragraphs:
        replace_in_runs(para.runs, ctx, test_log)

    # 替換表格儲存格中的佔位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs, ctx, test_log)

    # 替換頁眉和頁腳中的佔位符
    for sect in doc.sections:
        for para in sect.header.paragraphs:
            replace_in_runs(para.runs, ctx, test_log)
        for para in sect.footer.paragraphs:
            replace_in_runs(para.runs, ctx, test_log)

    doc.save(dst)
    link = dst
    return f"*[{params.get('desc', 'Word 檔案下載')}]({link})*\n\n" + "\n".join(test_log)





"""

def fill_word_with_context(params: dict) -> str:
    src, dst, ctx = params["filepath"], params["outputpath"], params["data_map"]
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    shutil.copy(src, dst)                             # 複製範本 → 新檔:contentReference[oaicite:5]{index=5}
    doc = Document(dst)

    test_log = "test_log : "

    def _replace_in_runs(runs):
        nonlocal test_log
        for run in runs:
            print("run.text:", run.text)
            test_log = test_log + "run.text:" + run.text + "\n"
            if "[" not in run.text:
                continue
            new_text = PATTERN.sub(lambda m: str(ctx.get(m.group(1), m.group(0))), run.text)
            print("new_text:", new_text)
            test_log = test_log + "new_text:" + new_text + "\n"
            print("--------------------------------")
            run.text = new_text

    # ── 1) 逐段落
    for p in doc.paragraphs:
        _replace_in_runs(p.runs)

    # ── 2) 逐表格儲存格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    _replace_in_runs(cp.runs)         # 表格中文字替換:contentReference[oaicite:6]{index=6}

    # ── 3) 頁眉／頁腳
    for sect in doc.sections:
        for p in sect.header.paragraphs:
            _replace_in_runs(p.runs)                  # 讀取 header API:contentReference[oaicite:7]{index=7}
        for p in sect.footer.paragraphs:
            _replace_in_runs(p.runs)

    doc.save(dst)
    #link = "https://office-mcp-dl.unieai.com/unieai-mcp-word/" + os.path.basename(dst)
    link = dst
    #response = "*[報價單A檔案下載](" + link_url + ")*"
    return f"*[{params.get('desc', 'Word 檔案下載')}]({link})* \n\n{test_log}"


"""



if __name__ == "__main__":
    main()
